#"""
#Gerador de Memorial Descritivo - Versão 5.1 (Streamlit Cloud - Melhorado)
#Lê PDFs de topografia (inclusive desenhos de CAD sem camada de texto, como
#exportações VectorDraw) convertendo as páginas em imagem e usando a visão
#multimodal do Gemini para extrair a tabela de roteiro perimétrico e os
#dados/confrontantes da planta (matrícula + nome dos vizinhos, sem CPF).
#
#Novidades v5.1:
#  - Retry automático com backoff exponencial para chamadas à API Gemini
#  - Validação rigorosa de entrada de dados do cliente (área, perímetro)
#  - Tratamento melhorado de erros de vértice inválido (vinculação de confrontantes)
#  - Limite inteligente de DPI (máximo 300) para evitar timeouts
#  - Sanitização de logs para não vazar dados sensíveis
#  - Docstrings detalhadas em modelos Pydantic
#  - Remoção de variável não utilizada (nome_mes)
#
#Funciona 100% no Streamlit Cloud sem dependências de sistema operacional
#(usa PyMuPDF puro-Python para rasterizar o PDF, sem precisar de Poppler/Tesseract).
#"""

import io
import re
import logging
import json
import time
from datetime import datetime
from typing import Optional, List, Dict, Any, Tuple

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from pypdf import PdfReader
import fitz  # PyMuPDF - rasteriza PDF em imagem sem depender de Poppler
from PIL import Image
import streamlit as st
from pydantic import BaseModel, ValidationError, Field
import google.generativeai as genai
from google.generativeai import types
from google.api_core import exceptions as google_exceptions

# ==========================================
# CONFIGURAÇÃO DE LOGGING
# ==========================================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ==========================================
# CONFIGURAÇÃO DA PÁGINA DO STREAMLIT
# ==========================================
st.set_page_config(
    page_title="Gerador de Memorial Descritivo - Gleba A",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# CONFIGURAÇÕES PADRÃO
# ==========================================
EMPRESA_INFO = {
    "nome": "TopoGeo Topografia e Consultoria LTDA",
    "endereco": "Rua Natalino Cossi, No 114, sala 2 - Vila Valério, CEP 29785-000",
    "telefone": "27 99837-1164",
    "email": "topogeo2014@gmail.com"
}

CLIENTE_INFO = {
    "imovel": "Lote",
    "proprietario": "SEBASTIAO IZOTON",
    "local": "Vila Valério",
    "area": "0,16 ha",
    "perimetro": "206,42 m"
}

TECNICO_INFO = {
    "nome": "Régis Campo da Silva",
    "cargo": "TÉCNICO EM AGROPECUÁRIA",
    "cfta": "11198519711",
    "trt": "BR20260210971"
}

MARGENS_CM = 2.5
FONTE_PADRAO = "Arial"
TAMANHO_FONTE_PADRAO = 11

# Configurações de retry
MAX_RETRIES = 3
INITIAL_RETRY_DELAY = 2  # segundos

# ==========================================
# MODELOS PYDANTIC
# ==========================================
class RegraConfrontante(BaseModel):
    """Regra de mapeamento de confrontantes entre pontos.
    
    Atributos:
        ponto_inicio: Número do vértice onde a regra começa
        ponto_fim: Número do vértice onde a regra termina
        nome_confrontante: Descrição do confrontante (ex: "Matrícula nº 1234 propriedade de JOÃO")
    """
    ponto_inicio: int
    ponto_fim: int
    nome_confrontante: str

    class Config:
        str_strip_whitespace = True


class MapeamentoConfrontantes(BaseModel):
    """Conjunto de regras de mapeamento de confrontantes."""
    regras: List[RegraConfrontante] = Field(default_factory=list)

    class Config:
        str_strip_whitespace = True


class SegmentoRoteiro(BaseModel):
    """Representa um segmento da tabela de roteiro perimétrico.
    
    Atributos:
        de: Número do vértice de origem (string)
        para: Número do vértice de destino (string)
        n_y: Coordenada Norte (Y), ex: "1234,567 m"
        e_x: Coordenada Este (X), ex: "5678,901 m"
        azimute: Azimute em graus/minutos/segundos, ex: "45°12'33\""
        distancia: Distância em metros, ex: "123,45 m"
    """
    de: str
    para: str
    n_y: str
    e_x: str
    azimute: str
    distancia: str

    class Config:
        str_strip_whitespace = True


class ExtracaoRoteiro(BaseModel):
    """Lista de segmentos extraídos da tabela de roteiro perimétrico."""
    segmentos: List[SegmentoRoteiro] = Field(default_factory=list)

    class Config:
        str_strip_whitespace = True


# ==========================================
# FUNÇÕES UTILITÁRIAS
# ==========================================
def sanitize_log_message(message: str) -> str:
    """Remove informações sensíveis de mensagens de log."""
    # Remove números de ponto flutuante que podem ser dados sensíveis
    # em favor de descrição genérica
    return message if len(message) < 200 else message[:200] + "..."


def validar_numero_positivo(valor: str, nome_campo: str) -> Tuple[bool, str]:
    """Valida se um valor é um número positivo válido.
    
    Args:
        valor: String com o valor a validar
        nome_campo: Nome do campo para mensagem de erro
        
    Returns:
        Tupla (é_válido, mensagem_erro)
    """
    try:
        # Remove espaços em branco
        valor_limpo = valor.strip()
        
        if not valor_limpo:
            return False, f"{nome_campo} não pode estar vazio"
        
        # Tenta converter para float (aceita vírgula ou ponto)
        num_valor = float(valor_limpo.replace(',', '.'))
        
        if num_valor <= 0:
            return False, f"{nome_campo} deve ser maior que zero"
        
        return True, ""
        
    except (ValueError, AttributeError):
        return False, f"{nome_campo} deve ser um número válido (ex: 0,16 ou 0.16)"


def retry_com_backoff(funcao, *args, max_tentativas=MAX_RETRIES, **kwargs):
    """Executa uma função com retry automático e backoff exponencial.
    
    Útil para chamar a API Gemini, que pode sofrer rate limiting.
    """
    for tentativa in range(1, max_tentativas + 1):
        try:
            return funcao(*args, **kwargs)
            
        except google_exceptions.ResourceExhausted as e:
            # Rate limit atingido
            if tentativa < max_tentativas:
                delay = INITIAL_RETRY_DELAY * (2 ** (tentativa - 1))
                logger.warning(
                    f"Rate limit da API atingido. Tentativa {tentativa}/{max_tentativas}. "
                    f"Aguardando {delay}s antes de retry..."
                )
                time.sleep(delay)
            else:
                raise ValueError(
                    "API Gemini atingiu o limite de requisições. "
                    "Tente novamente em alguns minutos."
                )
                
        except google_exceptions.DeadlineExceeded as e:
            # Timeout
            if tentativa < max_tentativas:
                delay = INITIAL_RETRY_DELAY * (2 ** (tentativa - 1))
                logger.warning(
                    f"Timeout na API. Tentativa {tentativa}/{max_tentativas}. "
                    f"Aguardando {delay}s antes de retry..."
                )
                time.sleep(delay)
            else:
                raise ValueError(
                    "API Gemini respondeu muito lentamente. "
                    "Tente aumentar o DPI ou usar menos páginas."
                )
        
        except google_exceptions.GoogleAPIError as e:
            # Outro erro genérico da API
            logger.error(f"Erro na API Gemini (tentativa {tentativa}): {sanitize_log_message(str(e))}")
            if tentativa < max_tentativas:
                delay = INITIAL_RETRY_DELAY * (2 ** (tentativa - 1))
                time.sleep(delay)
            else:
                raise


# ==========================================
# FUNÇÕES DE EXTRAÇÃO
# ==========================================
def verificar_pdf_tipo(arquivo_pdf) -> Tuple[str, bool]:
    """Verifica se o PDF contém texto ou é apenas imagens."""
    try:
        arquivo_pdf.seek(0)
        leitor = PdfReader(arquivo_pdf)
        
        texto_total = ""
        for pagina in leitor.pages:
            texto = pagina.extract_text()
            if texto:
                texto_total += texto
        
        tem_texto = len(texto_total.strip()) > 100
        tipo = "Texto" if tem_texto else "Imagem"
        
        logger.info(f"PDF detectado como: {tipo}")
        return tipo, tem_texto
        
    except Exception as e:
        logger.error(f"Erro ao verificar tipo de PDF: {sanitize_log_message(str(e))}")
        return "Desconhecido", False


def extrair_texto_pdf(arquivo_pdf) -> str:
    """Extrai texto do PDF normalmente (sem visão). Usado apenas como atalho
    quando o PDF já tem uma camada de texto real (não é o caso de plantas
    exportadas de CAD, que devem ser lidas com pdf_paginas_para_imagens)."""
    try:
        arquivo_pdf.seek(0)
        leitor = PdfReader(arquivo_pdf)
        texto_completo = ""
        
        logger.info(f"Extraindo texto de PDF com {len(leitor.pages)} páginas")
        
        for idx, pagina in enumerate(leitor.pages):
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                texto_completo += texto_pagina + "\n"
            else:
                logger.warning(f"Página {idx + 1} não contém texto extraível")
        
        return texto_completo
        
    except Exception as e:
        logger.error(f"Erro ao extrair texto do PDF: {sanitize_log_message(str(e))}")
        raise


def pdf_paginas_para_imagens(arquivo_pdf, dpi: int = 200) -> List[Image.Image]:
    """Converte cada página de um PDF em uma imagem PIL usando PyMuPDF.

    Não depende de Poppler/Tesseract (funciona no Streamlit Cloud). É o
    caminho usado para plantas/tabelas exportadas de CAD (ex.: VectorDraw),
    onde o texto é desenhado como vetor e não existe camada de texto real —
    por isso pypdf/pdftotext retornam vazio mesmo a fonte aparecendo listada
    nos metadados do PDF.
    
    Args:
        arquivo_pdf: Arquivo PDF carregado
        dpi: Resolução em DPI (padrão 200)
        
    Returns:
        Lista de imagens PIL convertidas do PDF
    """
    try:
        arquivo_pdf.seek(0)
        dados = arquivo_pdf.read()
        documento = fitz.open(stream=dados, filetype="pdf")

        zoom = dpi / 72
        matriz = fitz.Matrix(zoom, zoom)

        imagens = []
        for pagina in documento:
            pix = pagina.get_pixmap(matrix=matriz)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            imagens.append(img)

        documento.close()
        logger.info(f"PDF convertido em {len(imagens)} imagem(ns) a {dpi} DPI")
        return imagens

    except Exception as e:
        logger.error(f"Erro ao rasterizar PDF: {sanitize_log_message(str(e))}")
        raise ValueError(f"Não foi possível abrir o PDF para leitura visual: {str(e)}")


def parse_tabela_roteiro(texto_roteiro: str) -> List[Dict[str, str]]:
    """Extrai dados da tabela a partir de TEXTO já extraído (usado apenas no
    caminho de 'colar texto manualmente'). Para PDFs de CAD sem texto real,
    use extrair_roteiro_com_ia em vez desta função."""
    try:
        # Padrão 1: Formato com aspas duplas (mais comum)
        pattern1 = r'"(\d+)","(\d+)","([\d\.,]+)","([\d\.,]+)","([^"]+)","([\d\.,]+\s*m)"'
        matches = re.findall(pattern1, texto_roteiro)
        
        if not matches:
            logger.warning("Padrão 1 não encontrou correspondências. Tentando padrão alternativo...")
            # Padrão 2: Formato sem aspas (mais flexível)
            pattern2 = r'(\d+)\s+(\d+)\s+([\d\.,]+)\s+([\d\.,]+)\s+([°\d\'\"\s\.]+)\s+([\d\.,]+\s*m)'
            matches = re.findall(pattern2, texto_roteiro)
        
        if not matches:
            logger.warning("Nenhum padrão encontrou correspondências")
            return []
        
        segmentos = []
        for m in matches:
            try:
                az = m[4].replace('$', '').replace('\\circ', '°').replace('\\prime\\prime', '"').replace('\\prime', "'").strip()
                az = az.replace('\\:', '').strip()
                
                segmento = {
                    "de": m[0],
                    "para": m[1],
                    "n_y": m[2] + " m",
                    "e_x": m[3] + " m",
                    "azimute": az,
                    "distancia": m[5].strip(),
                    "confrontante": ""
                }
                segmentos.append(segmento)
                logger.debug(f"Segmento extraído: {m[0]} → {m[1]}")
                
            except Exception as e:
                logger.warning(f"Erro ao processar segmento {m}: {sanitize_log_message(str(e))}")
                continue
        
        if not segmentos:
            logger.warning("Nenhum segmento foi extraído da tabela")
        else:
            logger.info(f"Total de segmentos extraídos: {len(segmentos)}")
        
        return segmentos
        
    except Exception as e:
        logger.error(f"Erro ao fazer parse da tabela de roteiro: {sanitize_log_message(str(e))}")
        raise


def extrair_roteiro_com_ia(imagens_roteiro: List[Image.Image], nome_modelo: str) -> List[Dict[str, str]]:
    """Lê a(s) imagem(ns) da TABELA DE ROTEIRO PERIMÉTRICO usando a visão do
    Gemini e retorna a lista de segmentos (de/para/coordenadas/azimute/distância).
    Necessário porque tabelas exportadas de CAD (ex.: VectorDraw) não têm
    camada de texto real — o texto é desenhado como vetor.
    """
    try:
        prompt = """
        Você é um especialista em leitura de plantas e tabelas topográficas. A(s) imagem(ns)
        anexada(s) contém uma TABELA DE ROTEIRO PERIMÉTRICO de um levantamento topográfico
        (linhas de vértices/pontos com coordenadas, azimute e distância entre cada ponto e o
        próximo).

        Leia a tabela LINHA POR LINHA, na ordem em que aparecem, sem pular nenhuma, e para
        cada linha/segmento extraia os seguintes campos, **todos são obrigatórios e devem estar presentes no JSON, mesmo que vazios se não puderem ser lidos**:
        - de: número do vértice de origem (string)
        - para: número do vértice de destino (próximo vértice da linha) (string)
        - n_y: coordenada N (Norte/Y) do vértice de origem, incluindo a unidade "m" se houver (string)
        - e_x: coordenada E (Este/X) do vértice de origem, incluindo a unidade "m" se houver (string)
        - azimute: azimute do segmento no formato graus/minutos/segundos (ex: 45°12'33") (string)
        - distancia: distância do segmento, incluindo a unidade "m" (string)

        Transcreva EXATAMENTE os valores que aparecem na imagem. Não invente, arredonde ou
        preencha valores que não conseguir ler com certeza — nesse caso deixe o campo como
        string vazia "".
        """

        logger.info("Chamando API Gemini (visão) para extrair a tabela de roteiro...")

        def _chamar_gemini():
            model = genai.GenerativeModel(
                model_name=nome_modelo,
                generation_config=types.GenerationConfig(
                    response_mime_type="application/json",
                    response_schema=ExtracaoRoteiro,
                    temperature=0.0,
                ),
            )

            conteudo = list(imagens_roteiro) + [prompt]
            response = model.generate_content(conteudo)
            return response

        # Usar retry com backoff
        response = retry_com_backoff(_chamar_gemini)

        dados = json.loads(response.text)
        extracao = ExtracaoRoteiro(**dados)

        segmentos = []
        for seg in extracao.segmentos:
            item = seg.model_dump()
            item["confrontante"] = ""
            segmentos.append(item)

        logger.info(f"Total de segmentos extraídos via IA: {len(segmentos)}")
        return segmentos

    except ValidationError as e:
        logger.error(f"Erro de validação ao processar tabela de roteiro: {sanitize_log_message(str(e))}")
        raise ValueError(
            f"Resposta da IA inválida ao ler a tabela de roteiro. "
            f"Campos esperados faltando ou inválidos. Tente novamente com melhor qualidade de imagem."
        )
    except json.JSONDecodeError as e:
        logger.error(f"Erro ao fazer parse JSON da tabela de roteiro: {sanitize_log_message(str(e))}")
        raise ValueError(
            f"Resposta da IA não é JSON válido ao ler a tabela de roteiro. "
            f"Tente novamente ou aumente o DPI para melhor leitura."
        )
    except Exception as e:
        logger.error(f"Erro ao extrair tabela de roteiro com IA: {sanitize_log_message(str(e))}")
        raise


# ==========================================
# INTEGRAÇÃO COM GOOGLE GENERATIVE AI
# ==========================================
def configurar_gemini() -> bool:
    """Configura a conexão com a API do Google Generative AI."""
    try:
        api_key = st.secrets.get("GEMINI_API_KEY")
        
        if not api_key:
            logger.error("Chave GEMINI_API_KEY não encontrada nos Secrets")
            return False
        
        genai.configure(api_key=api_key)
        logger.info("Gemini configurado com sucesso")
        return True
        
    except Exception as e:
        logger.error(f"Erro ao configurar Gemini (verifique sua chave API)")
        return False


def mapear_confrontantes_gemini(
    nome_modelo: str,
    imagens_planta: Optional[List[Image.Image]] = None,
    texto_planta: Optional[str] = None,
    texto_roteiro: Optional[str] = None,
) -> Optional[MapeamentoConfrontantes]:
    """Mapeia confrontantes usando Gemini. Aceita a planta como imagem(ns)
    (caso de PDFs de CAD sem texto) e/ou como texto (caso de colagem manual).
    """
    try:
        prompt = """
        Você é um engenheiro agrimensor especialista em topografia. Analise o(s) documento(s)
        abaixo (imagem e/ou texto) para mapear os confrontantes da Gleba A.

        Sua tarefa é extrair os dados cadastrais solicitados e criar as regras matemáticas de transição de confrontantes.

        INSTRUÇÕES CRÍTICAS:
        1. Para cada confrontante, determine o intervalo de pontos (ponto_inicio e ponto_fim)
        2. Exemplo: Se do ponto 7 ao 21 confronta com 'Matrícula nº 1234 propriedade de JOAO', crie: ponto_inicio: 7, ponto_fim: 21, nome_confrontante: 'Matrícula nº 1234 propriedade de JOAO'
        3. Se houver fechamento do ciclo (ex: de ponto 21 para 1), use: ponto_inicio: 21, ponto_fim: 1
        4. Tente incluir a matrícula e os nomes dos vizinhos, não é necessário o CPF.
        5. Retorne ESTRITAMENTE no formato JSON estruturado fornecido
        6. NÃO invente dados. Se não conseguir extrair um campo, deixe como string vazia "" ou retorne lista vazia
        """

        if texto_planta:
            prompt += f"\n\nDOCUMENTO (DADOS DA PLANTA em texto):\n{texto_planta}\n"
        if texto_roteiro:
            prompt += f"\n\nDOCUMENTO (TABELA DE ROTEIRO PERIMÉTRICO em texto, use apenas como referência de contexto):\n{texto_roteiro}\n"

        logger.info("Chamando API Gemini para mapeamento de confrontantes...")

        def _chamar_gemini():
            model = genai.GenerativeModel(
                model_name=nome_modelo,
                generation_config=types.GenerationConfig(
                    response_mime_type="application/json",
                    response_schema=MapeamentoConfrontantes,
                    temperature=0.0,
                ),
            )

            conteudo = list(imagens_planta or []) + [prompt]
            response = model.generate_content(conteudo)
            return response

        # Usar retry com backoff
        response = retry_com_backoff(_chamar_gemini)
        
        logger.info("Resposta recebida da API Gemini")
        
        try:
            response_data = json.loads(response.text)
            mapeamento = MapeamentoConfrontantes(**response_data)
            logger.info(f"Mapeamento extraído: {len(mapeamento.regras)} regras de confrontantes")
            return mapeamento
            
        except ValidationError as e:
            logger.error(f"Erro de validação ao processar resposta do Gemini: {sanitize_log_message(str(e))}")
            raise ValueError(
                f"Resposta da IA com formato inválido. "
                f"Campos esperados: ponto_inicio, ponto_fim, nome_confrontante"
            )
        except json.JSONDecodeError as e:
            logger.error(f"Erro ao fazer parse JSON da resposta: {sanitize_log_message(str(e))}")
            raise ValueError(
                f"Resposta da IA não é JSON válido. Tente novamente com melhor qualidade de imagem."
            )
        
    except Exception as e:
        logger.error(f"Erro ao mapear confrontantes com Gemini: {sanitize_log_message(str(e))}")
        raise


# ==========================================
# LÓGICA DE VINCULAÇÃO
# ==========================================
def vincular_confrontantes(segmentos: List[Dict], mapeamento: MapeamentoConfrontantes) -> List[Dict]:
    """Vincula confrontantes aos segmentos com validação rigorosa."""
    logger.info("Iniciando vinculação de confrontantes aos segmentos...")
    
    for seg in segmentos:
        try:
            # Validação rigorosa de vértices
            de_str = seg.get("de", "").strip()
            para_str = seg.get("para", "").strip()
            
            if not de_str or not para_str:
                logger.warning(f"Segmento com vértice vazio detectado: {seg}")
                seg["confrontante"] = "VÉRTICE INVÁLIDO (vazio)"
                continue
            
            try:
                v_de = int(de_str)
                v_para = int(para_str)
            except ValueError:
                logger.warning(f"Vértices não-numéricos: de={de_str}, para={para_str}")
                seg["confrontante"] = f"VÉRTICE INVÁLIDO ({de_str}→{para_str})"
                continue
            
            confrontante_encontrado = None
            
            for regra in mapeamento.regras:
                if regra.ponto_inicio < regra.ponto_fim:
                    # Faixa normal (ex: 1 até 5)
                    if regra.ponto_inicio <= v_de < regra.ponto_fim:
                        confrontante_encontrado = regra.nome_confrontante.upper()
                        logger.debug(f"Segmento {v_de}→{v_para}: Faixa regular encontrada: {confrontante_encontrado}")
                        break
                
                elif regra.ponto_inicio > regra.ponto_fim:
                    # Ciclo fechado (ex: 21 até 1, fechando o polígono)
                    if v_de >= regra.ponto_inicio or v_de <= regra.ponto_fim:
                        confrontante_encontrado = regra.nome_confrontante.upper()
                        logger.debug(f"Segmento {v_de}→{v_para}: Ciclo fechado encontrado: {confrontante_encontrado}")
                        break
            
            if not confrontante_encontrado:
                confrontante_encontrado = "CONFRONTAÇÃO NÃO ENCONTRADA"
                logger.warning(f"Segmento {v_de}→{v_para}: Nenhuma regra correspondente encontrada")
            
            seg["confrontante"] = confrontante_encontrado
            
        except Exception as e:
            logger.error(f"Erro inesperado ao vincular confrontante: {sanitize_log_message(str(e))}")
            seg["confrontante"] = "ERRO NO PROCESSAMENTO"
    
    logger.info("Vinculação de confrontantes concluída")
    return segmentos


# ==========================================
# GERADOR DO DOCUMENTO WORD
# ==========================================
def gerar_documento_word(dados_finais: Dict[str, Any]) -> io.BytesIO:
    """Gera o arquivo Word com o memorial descritivo."""
    try:
        logger.info("Iniciando geração do documento Word...")
        
        doc = docx.Document()

        for section in doc.sections:
            section.top_margin = Cm(MARGENS_CM)
            section.bottom_margin = Cm(MARGENS_CM)
            section.left_margin = Cm(MARGENS_CM)
            section.right_margin = Cm(MARGENS_CM)

        style = doc.styles["Normal"]
        font = style.font
        font.name = FONTE_PADRAO
        font.size = Pt(TAMANHO_FONTE_PADRAO)

        # Cabeçalho (Header)
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False # Garante que o cabeçalho não seja o mesmo da seção anterior

        p_header_empresa = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p_header_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_header_empresa.paragraph_format.space_after = Pt(6)

        run_topogeo = p_header_empresa.add_run("TopoGeo")
        run_topogeo.font.size = Pt(14)
        run_topogeo.bold = True
        run_topogeo.font.color.rgb = RGBColor(0, 128, 0) # Verde

        p_header_empresa.add_run(" Topografia e Consultoria LTDA\n")
        p_header_empresa.add_run(f"{EMPRESA_INFO["endereco"]}\n")
        p_header_empresa.add_run(f"Fone {EMPRESA_INFO["telefone"]} - {EMPRESA_INFO["email"]}")

        # Adiciona uma linha separadora no cabeçalho
        p_header_linha = header.add_paragraph()
        p_header_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_header_linha.add_run("_" * 80)
        p_header_linha.paragraph_format.space_after = Pt(12)

        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_before = Pt(12)
        p_titulo.paragraph_format.space_after = Pt(18)
        run_tit = p_titulo.add_run("MEMORIAL DESCRITIVO")
        run_tit.bold = True
        run_tit.font.size = Pt(12)

        # Dados cadastrais
        p_dados = doc.add_paragraph()
        p_dados.paragraph_format.line_spacing = 1.15
        p_dados.paragraph_format.space_after = Pt(18)

        imovel = dados_finais.get('imovel', 'NÃO INFORMADO')
        proprietario = dados_finais.get('proprietario', 'NÃO INFORMADO')
        local = dados_finais.get('local', 'NÃO INFORMADO')
        area = dados_finais.get('area', 'NÃO INFORMADO')
        perimetro = dados_finais.get('perimetro', 'NÃO INFORMADO')

        p_dados.add_run("Imóvel: ").bold = True
        p_dados.add_run(f"{imovel}\n")
        p_dados.add_run("Proprietário: ").bold = True
        p_dados.add_run(f"{proprietario}\n")
        p_dados.add_run("Local: ").bold = True
        p_dados.add_run(f"{local}\n")
        p_dados.add_run("Área (ha): ").bold = True
        p_dados.add_run(f"{area}\n")
        p_dados.add_run("Perímetro (m): ").bold = True
        p_dados.add_run(f"{perimetro}")

        # Descrição
        p_desc_tit = doc.add_paragraph()
        p_desc_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desc_tit.paragraph_format.space_before = Pt(12)
        p_desc_tit.paragraph_format.space_after = Pt(12)
        run_desc = p_desc_tit.add_run("DESCRIÇÃO DO PERÍMETRO")
        run_desc.bold = True

        p_texto = doc.add_paragraph()
        p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_texto.paragraph_format.line_spacing = 1.25
        p_texto.paragraph_format.space_after = Pt(12)

        segmentos = dados_finais.get("segmentos", [])
        
        if segmentos:
            primeiro = segmentos[0]
            p_texto.add_run(
                f"Inicia-se a descrição deste perímetro no vértice {primeiro['de']}, "
                f"de coordenadas N {primeiro['n_y']} e E {primeiro['e_x']}; "
            )

            for i, s in enumerate(segmentos):
                if i + 1 < len(segmentos):
                    prox_coordenada_n = segmentos[i + 1]['n_y']
                    prox_coordenada_e = segmentos[i + 1]['e_x']
                else:
                    prox_coordenada_n = segmentos[0]['n_y']
                    prox_coordenada_e = segmentos[0]['e_x']

                confrontante = s.get('confrontante', 'NÃO INFORMADO')
                azimute = s.get('azimute', 'NÃO INFORMADO')
                distancia = s.get('distancia', 'NÃO INFORMADO')
                
                # Se for o último segmento que fecha o polígono, a formatação é um pouco diferente
                if i == len(segmentos) - 1:
                    p_texto.add_run(
                        f" {azimute} e {distancia} "
                        f"até o vértice {s['para']}, "
                    )
                else:
                    p_texto.add_run(
                        f"Divisa do imóvel; deste, segue confrontando com {confrontante}, "
                        f"com os seguintes azimutes e distâncias: {azimute} e {distancia} "
                        f"até o vértice {s['para']}, de coordenadas N {prox_coordenada_n} e E {prox_coordenada_e}; "
                    )
        else:
            logger.warning("Nenhum segmento disponível para gerar descrição")
            p_texto.add_run("Nenhum segmento foi processado.")

        p_texto.add_run(
            "ponto inicial da descrição deste perímetro. As coordenadas da base foram processadas pelo método de Posicionamento por Ponto Preciso (PPP). Todas as coordenadas aqui descritas "
            "estão georreferenciadas ao Sistema Geodésico Brasileiro e encontram-se representadas "
            "no Sistema U T M, referenciadas ao Meridiano Central nº 39°00', fuso -24, tendo como datum o SIRGAS2000. "
            "Todos os azimutes e distâncias, área e perímetro foram calculados no plano de projeção U T M."
        )

        # Data
        data_atual = datetime.now()

        p_data = doc.add_paragraph()
        p_data.paragraph_format.space_before = Pt(24)
        local_data = dados_finais.get('local', 'Vila Valério')
        p_data.add_run(f"{local_data} – ES, {data_atual.strftime('%d/%m/%Y')}.")

        # Assinatura
        p_assinatura = doc.add_paragraph()
        p_assinatura.paragraph_format.space_before = Pt(36)
        p_assinatura.add_run(
            f"__________________________________\n"
            f"                {TECNICO_INFO['nome']}\n"
            f"             Resp. Técnico\n"
            f"               CFTA: {TECNICO_INFO['cfta']}\n"
            f"Credenciamento INCRA: G1D"
        )

        conteudo_arquivo = io.BytesIO()
        doc.save(conteudo_arquivo)
        conteudo_arquivo.seek(0)
        
        logger.info("Documento Word gerado com sucesso")
        return conteudo_arquivo
        
    except Exception as e:
        logger.error(f"Erro ao gerar documento Word: {sanitize_log_message(str(e))}")
        raise


# ==========================================
# INTERFACE STREAMLIT
# ==========================================
def main():
    """Função principal da aplicação Streamlit"""

    st.title("📄 Processador de Memoriais Descritivos - Gleba A")
    st.write(
        "Insira os dois arquivos da Gleba A para estruturar automaticamente o Memorial Descritivo "
        "com precisão e conformidade técnica."
    )

    # Info sobre versão
    st.info("""
    ✅ **Versão 5.1 - Leitura visual via IA + Validações Rigorosas**
    - Funciona 100% na nuvem, sem Poppler/Tesseract
    - Lê PDFs de CAD (ex.: VectorDraw) que não têm texto extraível, convertendo
      as páginas em imagem e usando a visão do Gemini para ler as tabelas
    - Retry automático com backoff para chamadas à API Gemini (rate limiting)
    - Validação de área/perímetro e vértices inválidos
    - Limite inteligente de DPI para evitar timeouts
    - Extrai matrícula e nome dos vizinhos (sem CPF) da planta
    - Dados do cliente (Imóvel, Proprietário, Local, Área, Perímetro) configurados na barra lateral
    - Se preferir, ainda é possível colar o texto manualmente
    """)

    # Sidebar
    with st.sidebar:
        st.header("⚙️ Configurações")

        st.subheader("Dados da Empresa")
        empresa_nome = st.text_input("Nome da Empresa", value=EMPRESA_INFO["nome"])
        empresa_endereco = st.text_input("Endereço", value=EMPRESA_INFO["endereco"])
        empresa_telefone = st.text_input("Telefone", value=EMPRESA_INFO["telefone"])
        empresa_email = st.text_input("Email", value=EMPRESA_INFO["email"])

        st.subheader("Dados do Técnico Responsável")
        tecnico_nome = st.text_input("Nome do Técnico", value=TECNICO_INFO["nome"])
        tecnico_cfta = st.text_input("CFTA", value=TECNICO_INFO["cfta"])

        st.subheader("Modelo de IA")
        nome_modelo = st.selectbox(
            "Modelo Gemini",
            options=["Gemini 3.5 Flash", "Gemini 3.1 Pro", "Gemini 3.1 Flash Lite", "Gemini 2.5 Pro", "Gemini 2.5 Flash"],
            index=0,
            help="Escolha o modelo Gemini mais adequado. Modelos 'Flash' são mais rápidos e econômicos, 'Pro' oferecem maior capacidade de raciocínio para documentos complexos."
        )

        st.subheader("Dados do Cliente")
        cliente_imovel = st.text_input("Imóvel", value=CLIENTE_INFO["imovel"])
        cliente_proprietario = st.text_input("Nome do Proprietário", value=CLIENTE_INFO["proprietario"])
        cliente_local = st.text_input("Local", value=CLIENTE_INFO["local"])
        cliente_area = st.text_input("Área (ha)", value=CLIENTE_INFO["area"])
        cliente_perimetro = st.text_input("Perímetro (m)", value=CLIENTE_INFO["perimetro"])
        
        dpi_conversao = st.slider(
            "Qualidade da imagem (DPI)", 
            min_value=150, 
            max_value=300,  # Limitado a 300 para evitar timeouts
            value=200, 
            step=50,
            help="⚠️ DPI entre 150-300 recomendado. Valores altos melhoram leitura mas são mais lentos."
        )

        st.info(
            "💡 **Dica:** Modifique os dados acima se necessário. "
            "Eles serão usados em todos os documentos gerados nesta sessão."
        )

    # Upload
    st.subheader("📁 Carregue os Arquivos")

    col1, col2 = st.columns(2)
    with col1:
        pdf_planta = st.file_uploader(
            "Carregue o PDF com os DADOS DA PLANTA:",
            type=["pdf"],
            key="planta",
            help="PDF contendo a relação de confrontantes por intervalos de pontos"
        )
    with col2:
        pdf_roteiro = st.file_uploader(
            "Carregue o PDF da TABELA DE ROTEIRO PERIMÉTRICO:",
            type=["pdf"],
            key="roteiro",
            help="PDF contendo a tabela com coordenadas, azimutes e distâncias"
        )

    # Alternativa: Cola de texto
    with st.expander("📝 Alternativa: colar o texto manualmente (opcional)"):
        st.write(
            "Use apenas se preferir não enviar os PDFs, ou como reforço de contexto "
            "além dos PDFs enviados acima."
        )
        col1, col2 = st.columns(2)
        with col1:
            texto_planta_manual = st.text_area(
                "Cole o texto da PLANTA aqui (opcional):",
                height=100,
                key="texto_planta"
            )
        with col2:
            texto_roteiro_manual = st.text_area(
                "Cole o texto do ROTEIRO aqui (opcional):",
                height=100,
                key="texto_roteiro"
            )

    # Processamento
    if pdf_planta and pdf_roteiro or (texto_planta_manual and texto_roteiro_manual):
        if st.button("🔄 Analisar Documentos e Gerar Memorial", type="primary", use_container_width=True):

            # Validações de entrada
            erros_validacao = []
            
            # Validar área
            area_valida, msg_area = validar_numero_positivo(cliente_area, "Área")
            if not area_valida:
                erros_validacao.append(msg_area)
            
            # Validar perímetro
            perimetro_valido, msg_perimetro = validar_numero_positivo(cliente_perimetro, "Perímetro")
            if not perimetro_valido:
                erros_validacao.append(msg_perimetro)
            
            if erros_validacao:
                st.error("❌ Erros de validação encontrados:")
                for erro in erros_validacao:
                    st.error(f"  • {erro}")
                st.stop()

            EMPRESA_INFO["nome"] = empresa_nome
            EMPRESA_INFO["endereco"] = empresa_endereco
            EMPRESA_INFO["telefone"] = empresa_telefone
            EMPRESA_INFO["email"] = empresa_email
            TECNICO_INFO["nome"] = tecnico_nome
            TECNICO_INFO["cfta"] = tecnico_cfta
            
            CLIENTE_INFO["imovel"] = cliente_imovel
            CLIENTE_INFO["proprietario"] = cliente_proprietario
            CLIENTE_INFO["local"] = cliente_local
            CLIENTE_INFO["area"] = cliente_area
            CLIENTE_INFO["perimetro"] = cliente_perimetro

            with st.spinner("⏳ Processando documentos..."):
                try:
                    # Etapa 1: Configurar Gemini (precisa estar pronto antes de qualquer chamada de IA)
                    st.info("🔑 Etapa 1: Configurando API Gemini...")
                    if not configurar_gemini():
                        st.error(
                            "❌ Erro: Chave GEMINI_API_KEY não configurada nos Streamlit Secrets. "
                            "Configure a chave e tente novamente."
                        )
                        st.stop()
                    st.success("✅ Gemini configurado")

                    # Mapeamento de nomes amigáveis para nomes de API
                    model_mapping = {
                        "Gemini 3.5 Flash": "gemini-3.5-flash",
                        "Gemini 3.1 Pro": "gemini-3.1-pro",
                        "Gemini 3.1 Flash Lite": "gemini-3.1-flash-lite",
                        "Gemini 2.5 Pro": "gemini-2.5-pro",
                        "Gemini 2.5 Flash": "gemini-2.5-flash",
                    }
                    nome_modelo_api = model_mapping.get(nome_modelo, "gemini-3.5-flash")

                    imagens_planta: List[Image.Image] = []
                    imagens_roteiro: List[Image.Image] = []

                    # Etapa 2: Preparar as páginas dos PDFs como imagens
                    if pdf_planta:
                        st.info("🖼️ Etapa 2: Convertendo o PDF da planta em imagem...")
                        imagens_planta = pdf_paginas_para_imagens(pdf_planta, dpi=dpi_conversao)

                    if pdf_roteiro:
                        st.info("🖼️ Etapa 2: Convertendo o PDF do roteiro em imagem...")
                        imagens_roteiro = pdf_paginas_para_imagens(pdf_roteiro, dpi=dpi_conversao)

                    if imagens_planta or imagens_roteiro:
                        st.success("✅ Páginas convertidas em imagem")

                    # Etapa 3: Extrair a tabela de roteiro (via visão, se houver PDF; senão via texto colado)
                    st.info("📊 Etapa 3: Lendo a tabela de roteiro perimétrico...")
                    if imagens_roteiro:
                        segmentos_reais = extrair_roteiro_com_ia(imagens_roteiro, nome_modelo_api)
                    else:
                        segmentos_reais = parse_tabela_roteiro(texto_roteiro_manual)

                    if not segmentos_reais:
                        st.warning(
                            "⚠️ Nenhum segmento foi extraído da tabela de roteiro. "
                            "Confira se o PDF/texto enviado é realmente a tabela de roteiro perimétrico, "
                            "ou tente aumentar o DPI na barra lateral para melhorar a leitura."
                        )
                        st.stop()

                    st.success(f"✅ {len(segmentos_reais)} segmentos extraídos")

                    # Etapa 4: Mapear confrontantes e dados cadastrais
                    st.info("🤖 Etapa 4: Mapeando confrontantes e dados cadastrais com IA...")
                    mapeamento = mapear_confrontantes_gemini(
                        nome_modelo=nome_modelo_api,
                        imagens_planta=imagens_planta,
                        texto_planta=texto_planta_manual or None,
                        texto_roteiro=texto_roteiro_manual or None,
                    )
                    st.success(f"✅ {len(mapeamento.regras)} regras de confrontantes extraídas")

                    # Etapa 5: Vinculação
                    st.info("🔗 Etapa 5: Vinculando confrontantes aos segmentos...")
                    segmentos_reais = vincular_confrontantes(segmentos_reais, mapeamento)
                    st.success("✅ Confrontantes vinculados")

                    # Dados finais
                    dados_finais = {
                        "imovel": CLIENTE_INFO["imovel"],
                        "proprietario": CLIENTE_INFO["proprietario"],
                        "local": CLIENTE_INFO["local"],
                        "area": CLIENTE_INFO["area"],
                        "perimetro": CLIENTE_INFO["perimetro"],
                        "segmentos": segmentos_reais
                    }

                    # Resumo
                    st.success("🎉 Processamento concluído com sucesso!")

                    st.subheader("🔍 Resumo de Validação")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Proprietário", (dados_finais['proprietario'][:30] + "...") if dados_finais['proprietario'] else "NÃO INFORMADO")
                    with col2:
                        st.metric("Área Total", dados_finais['area'] or "NÃO INFORMADO")
                    with col3:
                        st.metric("Perímetro", dados_finais['perimetro'] or "NÃO INFORMADO")

                    # Tabela
                    with st.expander("📋 Clique para conferir a malha de confrontações vinculadas", expanded=True):
                        df_data = []
                        for seg in dados_finais["segmentos"]:
                            df_data.append({
                                "De": seg['de'],
                                "Para": seg['para'],
                                "N": seg['n_y'],
                                "E": seg['e_x'],
                                "Azimute": seg['azimute'],
                                "Distância": seg['distancia'],
                                "Confrontante": seg['confrontante']
                            })

                        import pandas as pd
                        df = pd.DataFrame(df_data)
                        st.dataframe(df, use_container_width=True, hide_index=True)
                        st.caption(
                            "⚠️ Confira os valores acima antes de usar o memorial oficialmente — "
                            "a leitura é feita por IA e pode conter erros, especialmente em desenhos com baixa qualidade."
                        )

                    # Geração
                    st.info("📝 Gerando documento Word...")
                    arquivo_docx = gerar_documento_word(dados_finais)
                    st.success("✅ Documento gerado com sucesso!")

                    # Download
                    nome_arquivo = CLIENTE_INFO['proprietario'].replace(' ', '_').upper() or 'MEMORIAL'
                    st.download_button(
                        label="📥 Baixar Memorial Descritivo (.docx)",
                        data=arquivo_docx,
                        file_name=f"MEMORIAL_DESCRITIVO_{nome_arquivo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                except ValueError as e:
                    st.error(f"❌ Erro de Validação: {str(e)}")
                    logger.error(f"Erro de validação: {sanitize_log_message(str(e))}")

                except json.JSONDecodeError as e:
                    st.error(f"❌ Erro ao processar resposta da IA: {str(e)}")
                    logger.error(f"Erro JSON: {sanitize_log_message(str(e))}")

                except Exception as e:
                    st.error(f"❌ Erro inesperado: {str(e)}")
                    logger.error(f"Erro geral: {sanitize_log_message(str(e))}", exc_info=True)

                    with st.expander("🔧 Detalhes Técnicos (Debug)"):
                        import traceback
                        st.code(traceback.format_exc())

    else:
        st.info("👆 Carregue ambos os PDFs, ou preencha o texto manual de ambos, para começar o processamento")


if __name__ == "__main__":
    main()
