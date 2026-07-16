# ==========================================
# ARQUIVO: gerador_anuencia_incra.py
# ==========================================
import io
import re
import json
import logging
import zipfile
from datetime import datetime
from typing import Dict, Any, List, Tuple
import streamlit as st
import google.generativeai as genai

try:
    import pypdf
except ImportError:
    pypdf = None

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

logger = logging.getLogger(__name__)


class GeradorAnuenciaIncraWord:
    def __init__(self, dados_empresa: Dict[str, str], dados_tecnico: Dict[str, str]):
        """
        Inicializa o gerador de anuência do INCRA com dados institucionais e do técnico.
        """
        self.dados_empresa = dados_empresa
        self.dados_tecnico = dados_tecnico

        # Configuração da API do Gemini obtida de forma segura dos secrets do Streamlit
        self.api_key = st.secrets.get("GEMINI_API_KEY")
        if self.api_key:
            genai.configure(api_key=self.api_key)

    def _estrutura_padrao(self) -> Dict[str, Any]:
        """
        Estrutura de fallback baseada fielmente nos modelos reais fornecidos.
        """
        return {
            "imovel_origem": {
                "proprietario": "AGOSTINHO IZOTON",
                "cpf": "215.894.707-10",
                "imovel": "GLEBA A",
                "localidade": "Vila Valerio-ES"
            },
            "confrontantes": [
                {
                    "confrontante_imovel": "Sitio Sete Quedas",
                    "confrontante_matricula": "8280",
                    "confrontante_comarca": "o Gabriel da Palha",
                    "confrontante_proprietario": "Elias Moro, Luiz Valentin Moro",
                    "confrontante_cpf": "780.485.677-68",
                    "vertices": [
                        {
                            "codigo": "G1D-P-06815",
                            "longitude": "40°17'14.01\"",
                            "latitude": "18°59'22.00\"",
                            "altitude": "58.39",
                            "vante": "G1D-P-06816",
                            "azimute": "02°15'",
                            "distancia": "41.66",
                            "confrontacao_completa": "CNS: 02.170-9 | Mat. 8280 | Sitio Sete Quedas; Elias Moro, Luiz Valentin Moro"
                        }
                    ]
                }
            ]
        }

    def _obter_dados_estruturados_com_ia(self, texto_memorial: str) -> Dict[str, Any]:
        """
        Usa o Gemini para analisar o memorial e extrair os dados estruturados.
        """
        estrutura_padrao = self._estrutura_padrao()

        if not self.api_key:
            logger.warning("Chave de API do Gemini não configurada nos secrets. Usando dados padrão.")
            return_data = estrutura_padrao
            return return_data

        prompt = f"""
        Você é um engenheiro cartógrafo especialista em georreferenciamento do INCRA.
        Sua tarefa é analisar o texto de um memorial descritivo ou relatório de cálculo/vértices e estruturar as informações do IMÓVEL DE ORIGEM e de TODOS os confrontantes identificados.

        TEXTO DO MEMORIAL DESCRITIVO:
        \"\"\"
        {texto_memorial}
        \"\"\"

        REGRAS CRÍTICAS DE FORMATAÇÃO DOS VALORES:
        - Para as coordenadas (longitude e latitude), você DEVE obrigatoriamente usar o símbolo de graus (°), minutos (') e segundos ("). Exemplo: 40°24'35.39"
        - Utilize sempre ponto (.) como separador decimal para altitude, distância e segundos das coordenadas.
        - Para Azimutes, garanta a presença do símbolo de grau (°). Exemplo: 292°21'

        Responda APENAS com o JSON estruturado abaixo, sem markdown ou textos explicativos:
        {{
            "imovel_origem": {{
                "proprietario": "NOME COMPLETO DO PROPRIETÁRIO PRINCIPAL",
                "cpf": "CPF DO PROPRIETÁRIO PRINCIPAL",
                "imovel": "NOME DO IMÓVEL RURAL DE ORIGEM",
                "localidade": "MUNICÍPIO OU LOCALIDADE DE ORIGEM"
            }},
            "confrontantes": [
                {{
                    "confrontante_imovel": "Nome do Imóvel Confrontante",
                    "confrontante_matricula": "Número da Matrícula/Transcrição",
                    "confrontante_comarca": "Nome da Comarca",
                    "confrontante_proprietario": "Nome do Proprietário Confrontante",
                    "confrontante_cpf": "CPF do Confrontante",
                    "vertices": [
                        {{
                            "codigo": "Código do Vértice",
                            "longitude": "Ex: 40°17'14.01\\"",
                            "latitude": "Ex: 18°59'22.00\\"",
                            "altitude": "Ex: 58.39",
                            "vante": "Código de Vante",
                            "azimute": "Ex: 292°21'",
                            "distancia": "Ex: 41.66",
                            "confrontacao_completa": "Descrição completa da confrontação"
                        }}
                    ]
                }}
            ]
        }}
        """

        try:
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(
                prompt,
                generation_config={"response_mime_type": "application/json"}
            )

            texto_resposta = response.text.strip()
            if texto_resposta.startswith("```json"):
                texto_resposta = texto_resposta.split("```json")[1].split("```")[0].strip()
            elif texto_resposta.startswith("```"):
                texto_resposta = texto_resposta.split("```")[1].split("```")[0].strip()

            dados = json.loads(texto_resposta)
            return dados
        except Exception as e:
            logger.error(f"Erro ao obter dados estruturados do Gemini: {str(e)}")
            return estrutura_padrao

    def _extrair_texto_memorial(self, conteudo_arquivo: bytes, nome_arquivo: str) -> str:
        texto_memorial = ""
        if nome_arquivo.lower().endswith(".pdf"):
            try:
                pdf_file = io.BytesIO(conteudo_arquivo)
                if pypdf:
                    reader = pypdf.PdfReader(pdf_file)
                    paginas_texto = []
                    for pagina in reader.pages:
                        texto_pag = pagina.extract_text()
                        if texto_pag:
                            paginas_texto.append(texto_pag)
                    texto_memorial = "\n".join(paginas_texto)
                else:
                    texto_memorial = conteudo_arquivo.decode("utf-8", errors="ignore")
            except Exception as pdf_err:
                logger.error(f"Falha ao extrair texto do PDF: {pdf_err}")
                texto_memorial = "Erro na leitura estruturada do PDF."
        elif nome_arquivo.lower().endswith(".docx"):
            try:
                doc_temp = Document(io.BytesIO(conteudo_arquivo))
                texto_memorial = "\n".join([p.text for p in doc_temp.paragraphs])
            except Exception as docx_err:
                logger.error(f"Não foi possível abrir o arquivo como .docx: {docx_err}")
                texto_memorial = conteudo_arquivo.decode("utf-8", errors="ignore")
        else:
            texto_memorial = conteudo_arquivo.decode("utf-8", errors="ignore")

        texto_memorial = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\xFF]', '', texto_memorial)
        return texto_memorial

    def gerar_documentos_pelo_memorial(
        self, conteudo_arquivo: bytes, nome_arquivo: str, dados_projeto: Dict[str, Any] = None
    ) -> List[Tuple[str, io.BytesIO]]:
        """
        Gera um documento Word separado para cada confrontante identificado no memorial.
        """
        if not conteudo_arquivo:
            raise ValueError("O conteúdo do arquivo está vazio ou corrompido.")

        texto_memorial = self._extrair_texto_memorial(conteudo_arquivo, nome_arquivo)
        dados_ia = self._obter_dados_estruturados_com_ia(texto_memorial)
        
        confrontantes = dados_ia.get("confrontantes") or self._estrutura_padrao()["confrontantes"]
        dados_origem = dados_ia.get("imovel_origem") or self._estrutura_padrao()["imovel_origem"]

        documentos: List[Tuple[str, io.BytesIO]] = []
        for dados_confrontante in confrontantes:
            nome_confrontante = str(
                dados_confrontante.get("confrontante_proprietario", "Confrontante")
            ).strip()
            buffer = self._montar_documento_confrontante(dados_confrontante, dados_origem)
            documentos.append((nome_confrontante, buffer))

        return documentos

    def gerar_documento_pelo_memorial(
        self, conteudo_arquivo: bytes, nome_arquivo: str, dados_projeto: Dict[str, Any] = None
    ) -> io.BytesIO:
        documentos = self.gerar_documentos_pelo_memorial(conteudo_arquivo, nome_arquivo, dados_projeto)
        return documentos[0][1]

    @staticmethod
    def gerar_zip_anuencias(documentos: List[Tuple[str, io.BytesIO]], prefixo_arquivo: str = "ANUENCIA_INCRA") -> io.BytesIO:
        zip_buffer = io.BytesIO()
        nomes_usados = set()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for nome_confrontante, buffer in documentos:
                nome_base = re.sub(r'[^A-Za-z0-9_\-]+', '_', nome_confrontante.strip().upper()) or "CONFRONTANTE"
                nome_arquivo = f"{prefixo_arquivo}_{nome_base}.docx"
                sufixo = 2
                nome_final = nome_arquivo
                while nome_final in nomes_usados:
                    nome_final = f"{prefixo_arquivo}_{nome_base}_{sufixo}.docx"
                    sufixo += 1
                nomes_usados.add(nome_final)

                buffer.seek(0)
                zip_file.writestr(nome_final, buffer.read())
                buffer.seek(0)

        zip_buffer.seek(0)
        return zip_buffer

    def _formatar_coordenada_estrita(self, valor: Any) -> str:
        """
        Preserva e corrige a formatação das coordenadas para manter o padrão de graus (°).
        Remove o sinal de menos (-) e garante o separador de ponto para os segundos decimais.
        """
        if not valor:
            return ""
        
        texto = str(valor).replace('-', '').strip()
        texto = texto.replace(',', '.')

        # Se a coordenada perdeu o símbolo de grau original, tenta reinserir baseando-se no padrão
        if "°" not in texto and len(texto) >= 4:
            match = re.match(r'^(\d{2})\s*([\d\']{2,})\s*([\d\.\"]+)', texto)
            if match:
                graus = match.group(1)
                minutos = match.group(2).replace("'", "")
                segundos = match.group(3).replace('"', '')
                texto = f"{graus}°{minutos}'{segundos}\""

        return texto

    def _formatar_azimute_estrito(self, valor: Any) -> str:
        """
        Garante que o azimute contenha o símbolo de graus (°) e separador de ponto.
        """
        if not valor:
            return ""
        texto = str(valor).strip().replace(",", ".")
        if "°" not in texto and len(texto) >= 2:
            match = re.match(r'^(\d{2,3})[\s\']?(\d{2})?', texto)
            if match:
                graus = match.group(1)
                minutos = match.group(2) if match.group(2) else "00"
                texto = f"{graus}°{minutos}'"
        return texto

    def _formatar_decimal_estrito_2_casas(self, valor: Any) -> str:
        """
        Força separador decimal por ponto (.) e garante exatamente 2 casas decimais.
        """
        if not valor:
            return "0.00"
        
        texto_limpo = re.sub(r'[^\d\.,]', '', str(valor).strip())
        texto_limpo = texto_limpo.replace(",", ".")
        
        try:
            num_float = float(texto_limpo)
            return f"{num_float:.2f}"
        except ValueError:
            return texto_limpo

    def _montar_documento_confrontante(
        self, dados_ia: Dict[str, Any], dados_origem: Dict[str, str]
    ) -> io.BytesIO:
        """
        Gera o documento Word da declaração de anuência do INCRA no formato paisagem (Landscape).
        Tabela configurada para manter os dados estritamente em uma única linha e sem quebras indesejadas.
        """
        doc = Document()

        # Configurar Orientação para PAISAGEM (Landscape)
        for section in doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.top_margin = Inches(0.65)
            section.bottom_margin = Inches(0.65)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # Estilo de Fonte Padrão Arial 11
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # 1. TÍTULO ORIGINAL
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after = Pt(10)
        run_titulo = p_titulo.add_run("DECLARA O DE RESPEITO DE LIMITES")
        run_titulo.bold = True
        run_titulo.font.size = Pt(12)
        run_titulo.font.name = 'Arial'

        # 2. DADOS DO PROPRIETÁRIO
        proprietario_origem = str(dados_origem.get("proprietario", "AGOSTINHO IZOTON")).upper()
        cpf_origem = str(dados_origem.get("cpf", "___.___.___-__"))
        localidade_origem = "Vila Valerio-ES"

        tecnico_nome = "Regis Campo da Silva"
        tecnico_cfta = "1119851971-1"
        codigo_incra = "G1D"

        p_corpo = doc.add_paragraph()
        p_corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_corpo.paragraph_format.space_after = Pt(10)
        p_corpo.paragraph_format.line_spacing = 1.15
        
        p_corpo.add_run("Eu, ").font.name = 'Arial'
        run_corp2 = p_corpo.add_run(f"{proprietario_origem}, CPF {cpf_origem}")
        run_corp2.bold = True
        run_corp2.font.name = 'Arial'
        
        p_corpo.add_run(f", residente em {localidade_origem}, e eu, ").font.name = 'Arial'
        run_corp4 = p_corpo.add_run(f"{tecnico_nome}")
        run_corp4.bold = True
        run_corp4.font.name = 'Arial'
        
        p_corpo.add_run(", Tecnico em Agropecuaria, CFTA ").font.name = 'Arial'
        run_cfta_val = p_corpo.add_run(f"{tecnico_cfta}")
        run_cfta_val.bold = True
        run_cfta_val.font.name = 'Arial'
        
        p_corpo.add_run(", credenciado pelo INCRA sob o codigo ").font.name = 'Arial'
        run_corp6 = p_corpo.add_run(f"{codigo_incra}")
        run_corp6.bold = True
        run_corp6.font.name = 'Arial'
        
        p_corpo.add_run(", declaramos sob as penas da Lei que quando dos trabalhos topograficos executados na citada propriedade ").font.name = 'Arial'
        run_corp8 = p_corpo.add_run("foram respeitados os limites de \"divisas in loco\"")
        run_corp8.bold = True
        run_corp8.font.name = 'Arial'
        
        p_corpo.add_run(" com os confrontantes abaixo relacionados, ").font.name = 'Arial'
        run_corp10 = p_corpo.add_run("não havendo qualquer litigio entre as partes.")
        run_corp10.bold = True
        run_corp10.font.name = 'Arial'

        # 3. CABEÇALHO CONFRONTANTES E DATA
        p_confrontantes_label = doc.add_paragraph()
        p_confrontantes_label.paragraph_format.space_after = Pt(4)
        run_conf_label = p_confrontantes_label.add_run(" Confrontantes:")
        run_conf_label.bold = True
        run_conf_label.font.name = 'Arial'

        data_atual = datetime.now()
        meses = [
            "JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO",
            "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO"
        ]
        texto_data = f"Vila Valerio-ES, {data_atual.day} de {meses[data_atual.month - 1]} de {data_atual.year}."

        p_data = doc.add_paragraph()
        p_data.paragraph_format.space_after = Pt(14)
        p_data.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_data.add_run(texto_data).font.name = 'Arial'

        # 4. TABELA
        tabela_vert = doc.add_table(rows=1, cols=8)
        tabela_vert.style = 'Table Grid'
        tabela_vert.autofit = False

        sub_headers = [
            "Código", "Longitude", "Latitude", "Altitude (m)",
            "Código", "Azimute", "Dist. (m)", "Confrontante"
        ]
        row_headers = tabela_vert.rows[0]
        for idx, text in enumerate(sub_headers):
            row_headers.cells[idx].text = text
            run_h = row_headers.cells[idx].paragraphs[0].runs[0]
            run_h.font.bold = True
            run_h.font.size = Pt(8.5)  # Cabeçalho ligeiramente menor
            run_h.font.name = 'Arial'

        # REDISTRIBUIÇÃO DE COLUNAS: Damos a largura máxima possível para a última coluna (Confrontante)
        # de forma que o texto longo caiba 100% horizontalmente.
        larguras_t2 = [
            Inches(0.9), Inches(1.15), Inches(1.15), Inches(0.75),
            Inches(0.9), Inches(0.65), Inches(0.65), Inches(3.85)  # Coluna Confrontante alargada de 2.3 para 3.85 polegadas
        ]

        # --- ADICIONAR DADOS DOS VÉRTICES ---
        vertices_dados = dados_ia.get("vertices", [])
        for v in vertices_dados:
            row = tabela_vert.add_row()
            cells = row.cells
            
            long_limpa = self._formatar_coordenada_estrita(v.get("longitude", ""))
            lat_limpa = self._formatar_coordenada_estrita(v.get("latitude", ""))
            alt_limpa = self._formatar_decimal_estrito_2_casas(v.get("altitude", ""))
            azimute_limpo = self._formatar_azimute_estrito(v.get("azimute", ""))
            dist_limpa = self._formatar_decimal_estrito_2_casas(v.get("distancia", ""))

            cells[0].text = str(v.get("codigo", ""))
            cells[1].text = long_limpa
            cells[2].text = lat_limpa
            cells[3].text = alt_limpa
            cells[4].text = str(v.get("vante", ""))
            cells[5].text = azimute_limpo
            cells[6].text = dist_limpa
            cells[7].text = str(v.get("confrontacao_completa", ""))

        # Formatação das Células para manter TUDO em 1 única linha
        for r_idx, row in enumerate(tabela_vert.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.width = larguras_t2[c_idx]
                p = cell.paragraphs[0]
                
                # Alinhamento à esquerda padrão
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Configurações de espaçamento super limpas (zeradas)
                p.paragraph_format.left_indent = None
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                
                if len(p.runs) > 0:
                    # REQUISITO CRÍTICO: Fonte reduzida para 7.5 para garantir que o texto do confrontante caiba em 1 única linha
                    p.runs[0].font.size = Pt(7.5)
                    p.runs[0].font.name = 'Arial'

        # Espaço proporcional antes das Assinaturas
        doc.add_paragraph().paragraph_format.space_before = Pt(32)

        # 5. TABELA DE ASSINATURAS HORIZONTAIS PARALELAS (Invisível)
        tabela_assinaturas = doc.add_table(rows=2, cols=2)
        tabela_assinaturas.autofit = False
        tabela_assinaturas.columns[0].width = Inches(4.7)
        tabela_assinaturas.columns[1].width = Inches(4.7)

        cells_as = tabela_assinaturas.rows[0].cells
        cells_as[0].text = "__________________________________________________"
        cells_as[1].text = "__________________________________________________"

        cells_nomes = tabela_assinaturas.rows[1].cells

        # Coluna Proprietário de Origem
        p_origem = cells_nomes[0].paragraphs[0]
        p_origem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_o1 = p_origem.add_run(f"\n{proprietario_origem}\n")
        run_o1.bold = True
        p_origem.add_run(f"CPF: {cpf_origem}")

        # Coluna Proprietário Confrontante (Nome limpo)
        p_confrontante = cells_nomes[1].paragraphs[0]
        p_confrontante.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nome_vizinho = str(dados_ia.get("confrontante_proprietario", "")).upper()
        
        run_v1 = p_confrontante.add_run(f"\n{nome_vizinho}\n")
        run_v1.bold = True

        for row in tabela_assinaturas.rows:
            for cell in row.cells:
                cell.width = Inches(4.7)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Arial'

        # 6. BLOCO DO RESPONSÁVEL TÉCNICO
        doc.add_paragraph().paragraph_format.space_before = Pt(22)
        p_linha_rt = doc.add_paragraph()
        p_linha_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_linha_rt.add_run("_____________________________________")

        p_info_rt = doc.add_paragraph()
        p_info_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_rt1 = p_info_rt.add_run(f"{tecnico_nome}\n")
        run_rt1.bold = True
        p_info_rt.add_run(f"CFTA: {tecnico_cfta}")

        for run in p_info_rt.runs:
            run.font.size = Pt(9.5)
            run.font.name = 'Arial'

        # 7. ANEXOS (Espaçamento simples)
        doc.add_paragraph().paragraph_format.space_before = Pt(18)
        p_anexos = doc.add_paragraph()
        p_anexos.add_run("Anexos:  ").bold = True
        p_anexos.add_run("Planta do Imóvel  ")
        p_anexos.add_run("Memorial Descritivo do Imóvel")
        for run in p_anexos.runs:
            run.font.size = Pt(9)
            run.font.name = 'Arial'

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
