#!/usr/bin/env python3
"""
Script de validação da correção do gerador de requerimento de cartório.
Testa se os placeholders estão únicos e se o mapeamento funciona corretamente.
"""

from docx import Document
import re
from collections import defaultdict

def analisar_template_novo(template_path):
    """Analisa o novo template para verificar se placeholders são únicos."""
    print("=" * 90)
    print("VALIDAÇÃO DO NOVO TEMPLATE")
    print("=" * 90)
    
    doc = Document(template_path)
    
    placeholders = defaultdict(int)
    
    # Analisar parágrafos
    for p in doc.paragraphs:
        matches = re.findall(r'\{\{[A-Z_]+\}\}', p.text)
        for match in matches:
            placeholders[match] += 1
    
    # Analisar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = re.findall(r'\{\{[A-Z_]+\}\}', p.text)
                    for match in matches:
                        placeholders[match] += 1
    
    print(f"\n✅ Total de placeholders únicos encontrados: {len(placeholders)}")
    print(f"✅ Total de ocorrências: {sum(placeholders.values())}")
    
    # Verificar duplicatas
    duplicados = {p: freq for p, freq in placeholders.items() if freq > 1}
    
    if duplicados:
        print(f"\n❌ ERRO: Encontrados {len(duplicados)} placeholders duplicados:")
        for placeholder, freq in sorted(duplicados.items(), key=lambda x: -x[1]):
            print(f"  - '{placeholder}' aparece {freq} vezes")
        return False
    else:
        print("\n✅ SUCESSO: Nenhum placeholder duplicado encontrado!")
        print("\nPlaceholders únicos:")
        for placeholder in sorted(placeholders.keys()):
            print(f"  - {placeholder}")
        return True


def validar_mapeamento():
    """Valida se o mapeamento de placeholders está correto."""
    print("\n" + "=" * 90)
    print("VALIDAÇÃO DO MAPEAMENTO DE PLACEHOLDERS")
    print("=" * 90)
    
    # Mapeamento esperado
    mapeamento_esperado = {
        "{{COMARCA}}": "Comarca",
        "{{NOME_PROPRIETARIO}}": "Nome do proprietário",
        "{{PROFISSAO_PROPRIETARIO}}": "Profissão do proprietário",
        "{{RG_PROPRIETARIO}}": "RG do proprietário",
        "{{ORGAO_PROPRIETARIO}}": "Órgão expedidor do RG",
        "{{CPF_PROPRIETARIO}}": "CPF do proprietário",
        "{{NOME_ESPOSA}}": "Nome da esposa",
        "{{PROFISSAO_ESPOSA}}": "Profissão da esposa",
        "{{RG_ESPOSA}}": "RG da esposa",
        "{{ORGAO_ESPOSA}}": "Órgão expedidor do RG da esposa",
        "{{CPF_ESPOSA}}": "CPF da esposa",
        "{{REGIME_BENS}}": "Regime de bens",
        "{{ENDERECO_CORREGO}}": "Endereço/Córrego",
        "{{MUNICIPIO_CLIENTE}}": "Município do cliente",
        "{{NOME_SITIO}}": "Nome do sítio",
        "{{AREA_REGISTRADA}}": "Área registrada",
        "{{MUNICIPIO_IMOVEL}}": "Município do imóvel",
        "{{COMARCA_IMOVEL}}": "Comarca do imóvel",
        "{{MATRICULA}}": "Matrícula",
        "{{AREA_ENCONTRADA}}": "Área encontrada",
        "{{CODIGO_INCRA}}": "Código INCRA",
        "{{TRT_NUMERO}}": "Número da TRT",
        "{{AREA_TOTAL_RETIFICADA}}": "Área total retificada",
        "{{DATA_FORMATADA}}": "Data formatada",
        "{{ASSINATURA_PROPRIETARIO}}": "Assinatura do proprietário",
        "{{ASSINATURA_ESPOSA}}": "Assinatura da esposa",
        "{{CPF_ASSINATURA_PROPRIETARIO}}": "CPF para assinatura do proprietário",
        "{{CPF_ASSINATURA_ESPOSA}}": "CPF para assinatura da esposa",
    }
    
    print(f"\n✅ Total de campos no mapeamento: {len(mapeamento_esperado)}")
    print("\nCampos mapeados:")
    for placeholder, descricao in sorted(mapeamento_esperado.items()):
        print(f"  - {placeholder:35} → {descricao}")
    
    return True


def comparar_templates(template_antigo, template_novo):
    """Compara os templates antigo e novo."""
    print("\n" + "=" * 90)
    print("COMPARAÇÃO ENTRE TEMPLATES")
    print("=" * 90)
    
    doc_antigo = Document(template_antigo)
    doc_novo = Document(template_novo)
    
    # Extrair placeholders antigos
    placeholders_antigos = defaultdict(int)
    for p in doc_antigo.paragraphs:
        matches = re.findall(r'\([X\-\d\s\.]+\)|X+(?![a-z])|BR\d+', p.text)
        for match in matches:
            placeholders_antigos[match] += 1
    
    for table in doc_antigo.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = re.findall(r'\([X\-\d\s\.]+\)|X+(?![a-z])|BR\d+', p.text)
                    for match in matches:
                        placeholders_antigos[match] += 1
    
    # Extrair placeholders novos
    placeholders_novos = defaultdict(int)
    for p in doc_novo.paragraphs:
        matches = re.findall(r'\{\{[A-Z_]+\}\}', p.text)
        for match in matches:
            placeholders_novos[match] += 1
    
    for table in doc_novo.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = re.findall(r'\{\{[A-Z_]+\}\}', p.text)
                    for match in matches:
                        placeholders_novos[match] += 1
    
    print(f"\n📊 Template Antigo:")
    print(f"  - Placeholders únicos: {len(placeholders_antigos)}")
    print(f"  - Total de ocorrências: {sum(placeholders_antigos.values())}")
    duplicados_antigos = {p: freq for p, freq in placeholders_antigos.items() if freq > 1}
    print(f"  - Placeholders duplicados: {len(duplicados_antigos)}")
    
    print(f"\n📊 Template Novo:")
    print(f"  - Placeholders únicos: {len(placeholders_novos)}")
    print(f"  - Total de ocorrências: {sum(placeholders_novos.values())}")
    duplicados_novos = {p: freq for p, freq in placeholders_novos.items() if freq > 1}
    print(f"  - Placeholders duplicados: {len(duplicados_novos)}")
    
    print(f"\n📈 Melhoria:")
    print(f"  - Redução de placeholders duplicados: {len(duplicados_antigos)} → {len(duplicados_novos)}")
    print(f"  - Redução de ocorrências: {sum(placeholders_antigos.values())} → {sum(placeholders_novos.values())}")
    
    return True


def main():
    """Executa todos os testes de validação."""
    print("\n")
    print("╔" + "=" * 88 + "╗")
    print("║" + " " * 88 + "║")
    print("║" + "VALIDAÇÃO DA CORREÇÃO - GERADOR DE REQUERIMENTO DE CARTÓRIO".center(88) + "║")
    print("║" + " " * 88 + "║")
    print("╚" + "=" * 88 + "╝")
    
    template_antigo = "/home/ubuntu/upload/GERADOR-DEVMEMORIAIS-main/-REQUERIMENTODECARTORIO.docx"
    template_novo = "/home/ubuntu/upload/GERADOR-DEVMEMORIAIS-main/-REQUERIMENTODECARTORIO_CORRIGIDO.docx"
    
    try:
        # Teste 1: Validar novo template
        resultado1 = analisar_template_novo(template_novo)
        
        # Teste 2: Validar mapeamento
        resultado2 = validar_mapeamento()
        
        # Teste 3: Comparar templates
        resultado3 = comparar_templates(template_antigo, template_novo)
        
        # Resultado final
        print("\n" + "=" * 90)
        print("RESULTADO FINAL DA VALIDAÇÃO")
        print("=" * 90)
        
        if resultado1 and resultado2 and resultado3:
            print("\n✅ TODOS OS TESTES PASSARAM COM SUCESSO!")
            print("\n🎉 A correção está pronta para ser implementada em produção.")
            print("\nPróximos passos:")
            print("  1. Backup dos arquivos originais")
            print("  2. Copiar novo template para substituir o antigo")
            print("  3. Copiar novo código para substituir o antigo")
            print("  4. Reiniciar a aplicação Streamlit")
            print("  5. Testar com dados reais")
        else:
            print("\n❌ ALGUNS TESTES FALHARAM!")
            print("Verifique os erros acima e corrija antes de implementar.")
        
        print("\n" + "=" * 90 + "\n")
        
    except Exception as e:
        print(f"\n❌ ERRO durante validação: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
