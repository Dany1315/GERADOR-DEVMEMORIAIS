# ============================================================
# GERADOR DE MEMORIAL DESCRITIVO EM WORD
# Formato: Parágrafo Contínuo com Cabeçalho da Empresa
# Versão Corrigida: Sem duplicação de cabeçalho e matrícula
# ============================================================

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from typing import Dict, List, Any


class GeradorMemorialWord:
    """
    Gera memorial descritivo em formato Word com:
    - Cabeçalho da empresa (sem duplicação)
    - Dados do imóvel
    - Descrição contínua (um parágrafo único)
    - Assinatura do técnico com CPF
    """

    def __init__(self, dados_empresa: Dict[str, str], dados_tecnico: Dict[str, str]):
        """
        Inicializa o gerador com dados da empresa e técnico.
        
        Args:
            dados_empresa: Dict com 'nome', 'endereco', 'telefone', 'email'
            dados_tecnico: Dict com 'nome', 'cfta', 'cpf', 'trt'
        """
        self.dados_empresa = dados_empresa
        self.dados_tecnico = dados_tecnico

    def _adicionar_linha_separadora(self, doc: Document):
        """Adiciona uma linha separadora no documento."""
        paragraph = doc.add_paragraph()
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _formatar_nome_empresa(self, doc: Document, nome: str):
        """Adiciona o nome da empresa em verde e centralizado."""
        p = doc.add_paragraph(nome)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0)  # Verde

    def _formatar_texto_centralizado(self, doc: Document, texto: str, tamanho: int = 11, negrito: bool = False):
        """Adiciona texto centralizado ao documento."""
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(tamanho)
        if negrito:
            run.font.bold = True

    def _adicionar_cabecalho_empresa(self, doc: Document):
        """Adiciona o cabeçalho com dados da empresa (SEM DUPLICAÇÃO)."""
        # Nome da empresa (em verde) - APENAS "TopoGeo"
        self._formatar_nome_empresa(doc, 'TopoGeo')
        
        # Tipo de empresa
        self._formatar_texto_centralizado(doc, 'Topografia e Consultoria LTDA', tamanho=11)
        
        # Endereço + Telefone em uma linha
        endereco = self.dados_empresa.get('endereco', '')
        telefone = self.dados_empresa.get('telefone', '')
        endereco_telefone = f'{endereco} Fone {telefone}'
        self._formatar_texto_centralizado(doc, endereco_telefone, tamanho=11)
        
        # Email
        email = self.dados_empresa.get('email', '')
        self._formatar_texto_centralizado(doc, email, tamanho=11)
        
        # Linha separadora
        self._adicionar_linha_separadora(doc)

    def _adicionar_titulo(self, doc: Document):
        """Adiciona o título 'MEMORIAL DESCRITIVO' centralizado e negrito."""
        p = doc.add_paragraph('MEMORIAL DESCRITIVO')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(12)
        run.font.bold = True

    def _adicionar_dados_imovel(self, doc: Document, dados: Dict[str, Any]):
        """Adiciona os dados do imóvel (proprietário, município, etc)."""
        # Proprietário
        proprietario = dados.get('proprietario', 'N/A')
        p = doc.add_paragraph(f'Proprietário: {proprietario}')
        p.paragraph_format.left_indent = Inches(0)
        
        # Município
        municipio = dados.get('municipio', 'N/A')
        p = doc.add_paragraph(f'Município: {municipio}')
        
        # Comarca
        comarca = dados.get('comarca', 'N/A')
        p = doc.add_paragraph(f'Comarca: {comarca}')
        
        # TRT
        trt = dados.get('trt', 'N/A')
        p = doc.add_paragraph(f'TRT: {trt}')
        
        # Perímetro
        perimetro = dados.get('perimetro', 'N/A')
        p = doc.add_paragraph(f'Perímetro: {perimetro} m')
        
        # Área
        area = dados.get('area', 'N/A')
        p = doc.add_paragraph(f'Area: {area} m²')
        
        # Matrícula
        matricula = dados.get('matricula', 'N/A')
        p = doc.add_paragraph(f'MAT. {matricula}')

    def _construir_descricao_continua(self, segmentos: List[Dict[str, Any]]) -> str:
        """
        Constrói a descrição como um parágrafo único e contínuo.
        
        Args:
            segmentos: Lista de dicionários com dados dos segmentos
            
        Returns:
            String com a descrição completa e contínua
        """
        if not segmentos:
            return "Nenhum segmento disponível."
        
        # Inicia com o primeiro segmento
        primeiro = segmentos[0]
        confrontante_inicial = primeiro.get('confrontante', 'N/A')
        matricula_inicial = primeiro.get('matricula', 'N/A')
        coord_y_1 = primeiro.get('coord_y', 'N/A')
        coord_x_1 = primeiro.get('coord_x', 'N/A')
        
        # Começa a descrição
        descricao = (
            f"Inicia-se a descrição deste perímetro no vértice 1, de coordenadas N(Y) {coord_y_1} m e "
            f"E(X) {coord_x_1} m, situado na divisa com {confrontante_inicial} (MAT. {matricula_inicial}); "
            f"deste, segue confrontando com {confrontante_inicial} (MAT. {matricula_inicial}), "
            f"com o(s) seguinte(s) azimute(s) e distância(s): "
        )
        
        # Adiciona cada segmento
        for idx, segmento in enumerate(segmentos):
            azimute = segmento.get('azimute', 'N/A')
            distancia = segmento.get('distancia', 'N/A')
            vértice_para = idx + 2  # Próximo vértice
            coord_y_para = segmento.get('coord_y_para', 'N/A')
            coord_x_para = segmento.get('coord_x_para', 'N/A')
            confrontante = segmento.get('confrontante', 'N/A')
            matricula = segmento.get('matricula', 'N/A')
            
            # Adiciona o segmento
            if idx < len(segmentos) - 1:
                # Não é o último segmento
                descricao += (
                    f"{azimute} e {distancia} m até o vértice {vértice_para}, de coordenadas "
                    f"N(Y) {coord_y_para} m e E(X) {coord_x_para} m; "
                )
                
                # Verifica se o próximo segmento tem confrontante diferente
                proximo_segmento = segmentos[idx + 1]
                proximo_confrontante = proximo_segmento.get('confrontante', 'N/A')
                proximo_matricula = proximo_segmento.get('matricula', 'N/A')
                
                if proximo_confrontante != confrontante or proximo_matricula != matricula:
                    descricao += (
                        f"deste, segue confrontando com {proximo_confrontante} (MAT. {proximo_matricula}), "
                        f"com o(s) seguinte(s) azimute(s) e distância(s): "
                    )
            else:
                # Último segmento - fecha o perímetro
                descricao += (
                    f"{azimute} e {distancia} m até o vértice 1, ponto inicial da descrição deste perímetro."
                )
        
        return descricao

    def _adicionar_secao_descricao(self, doc: Document):
        """Adiciona a seção 'DESCRIÇÃO' centralizada e negrita."""
        p = doc.add_paragraph('DESCRIÇÃO')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(11)
        run.font.bold = True

    def _adicionar_descricao_continua(self, doc: Document, texto_descricao: str):
        """Adiciona a descrição como um parágrafo único e contínuo."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Adiciona o texto
        run = p.add_run(texto_descricao)
        run.font.size = Pt(11)
        
        # Espaçamento entre linhas
        p.paragraph_format.line_spacing = 1.5

    def _adicionar_data_local(self, doc: Document, local: str, data: str):
        """Adiciona data e local no final do documento."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{local}, {data}')
        run.font.size = Pt(11)

    def _adicionar_assinatura(self, doc: Document):
        """Adiciona a assinatura do técnico com CPF."""
        # Espaço em branco
        doc.add_paragraph()
        
        # Linha de assinatura
        p = doc.add_paragraph('_____________________________')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nome do técnico
        nome_tecnico = self.dados_tecnico.get('nome', 'N/A')
        p = doc.add_paragraph(nome_tecnico)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(11)
        
        # Profissão
        p = doc.add_paragraph('Técnico em Agropecuária')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(11)
        
        # CFTA, CPF e TRT
        cfta = self.dados_tecnico.get('cfta', 'N/A')
        cpf = self.dados_tecnico.get('cpf', 'N/A')
        trt = self.dados_tecnico.get('trt', 'N/A')
        
        p = doc.add_paragraph(f'CFTA: {cfta} | CPF: {cpf} | TRT: {trt}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(11)

    def gerar_documento(self, dados_finais: Dict[str, Any]) -> bytes:
        """
        Gera o documento Word completo com memorial descritivo.
        
        Args:
            dados_finais: Dicionário com todos os dados necessários
            
        Returns:
            Bytes do documento Word gerado
        """
        # Cria novo documento
        doc = Document()
        
        # Define margens padrão
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 1. Adiciona cabeçalho da empresa
        self._adicionar_cabecalho_empresa(doc)
        
        # 2. Adiciona espaço
        doc.add_paragraph()
        
        # 3. Adiciona título
        self._adicionar_titulo(doc)
        
        # 4. Adiciona espaço
        doc.add_paragraph()
        
        # 5. Adiciona dados do imóvel
        self._adicionar_dados_imovel(doc, dados_finais)
        
        # 6. Adiciona espaço
        doc.add_paragraph()
        
        # 7. Adiciona seção DESCRIÇÃO
        self._adicionar_secao_descricao(doc)
        
        # 8. Adiciona espaço
        doc.add_paragraph()
        
        # 9. Constrói e adiciona descrição contínua
        segmentos = dados_finais.get('segmentos', [])
        texto_descricao = self._construir_descricao_continua(segmentos)
        self._adicionar_descricao_continua(doc, texto_descricao)
        
        # 10. Adiciona espaço
        doc.add_paragraph()
        
        # 11. Adiciona data e local
        local = dados_finais.get('municipio', 'Local')
        data = dados_finais.get('data', '21 de julho de 2026')
        self._adicionar_data_local(doc, local, data)
        
        # 12. Adiciona assinatura
        self._adicionar_assinatura(doc)
        
        # Salva em BytesIO e retorna
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
