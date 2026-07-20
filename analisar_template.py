#!/usr/bin/env python3
"""Analisa os placeholders presentes no template de requerimento de cartório."""

from docx import Document
import re
from collections import defaultdict

doc = Document("/home/ubuntu/upload/GERADOR-DEVMEMORIAIS-main/-REQUERIMENTODECARTORIO.docx")

print("=" * 90)
print("ANÁLISE DE PLACEHOLDERS NO TEMPLATE DE REQUERIMENTO DE CARTÓRIO")
print("=" * 90)

placeholders_por_posicao = []
placeholders_unicos = defaultdict(int)

# Analisar parágrafos
print("\n📄 PARÁGRAFOS COM PLACEHOLDERS:")
print("-" * 90)

for para_idx, p in enumerate(doc.paragraphs):
    texto = p.text
    # Encontra todos os placeholders
    matches = list(re.finditer(r'\([X\-\d\s\.]+\)|X+(?![a-z])|BR\d+', texto))
    
    if matches:
        print(f"\n[Parágrafo {para_idx}]")
        print(f"Texto: {texto[:120]}{'...' if len(texto) > 120 else ''}")
        print(f"Placeholders encontrados:")
        
        for match in matches:
            placeholder = match.group()
            placeholders_unicos[placeholder] += 1
            placeholders_por_posicao.append({
                'tipo': 'parágrafo',
                'índice': para_idx,
                'placeholder': placeholder,
                'contexto': texto[max(0, match.start()-30):min(len(texto), match.end()+30)]
            })
            print(f"  - '{placeholder}' em posição {match.start()}-{match.end()}")

# Analisar tabelas
print("\n\n📊 TABELAS COM PLACEHOLDERS:")
print("-" * 90)

for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para_idx, p in enumerate(cell.paragraphs):
                texto = p.text
                matches = list(re.finditer(r'\([X\-\d\s\.]+\)|X+(?![a-z])|BR\d+', texto))
                
                if matches:
                    print(f"\n[Tabela {table_idx}, Linha {row_idx}, Coluna {cell_idx}]")
                    print(f"Texto: {texto[:120]}{'...' if len(texto) > 120 else ''}")
                    print(f"Placeholders encontrados:")
                    
                    for match in matches:
                        placeholder = match.group()
                        placeholders_unicos[placeholder] += 1
                        placeholders_por_posicao.append({
                            'tipo': 'tabela',
                            'índice': f"{table_idx},{row_idx},{cell_idx}",
                            'placeholder': placeholder,
                            'contexto': texto[max(0, match.start()-30):min(len(texto), match.end()+30)]
                        })
                        print(f"  - '{placeholder}' em posição {match.start()}-{match.end()}")

# Resumo
print("\n\n" + "=" * 90)
print("RESUMO DE PLACEHOLDERS ÚNICOS")
print("=" * 90)

print(f"\nTotal de placeholders únicos: {len(placeholders_unicos)}")
print(f"Total de ocorrências: {sum(placeholders_unicos.values())}")

print("\nPlaceholders e suas frequências:")
print("-" * 90)

for placeholder in sorted(placeholders_unicos.keys()):
    freq = placeholders_unicos[placeholder]
    print(f"  '{placeholder:30}' → {freq:2} ocorrência(s)")

# Identificar duplicatas
print("\n" + "=" * 90)
print("⚠️  PLACEHOLDERS DUPLICADOS (PROBLEMA CRÍTICO!)")
print("=" * 90)

duplicados = {p: freq for p, freq in placeholders_unicos.items() if freq > 1}

if duplicados:
    print(f"\n⛔ Encontrados {len(duplicados)} placeholders que aparecem mais de uma vez:\n")
    for placeholder in sorted(duplicados.keys(), key=lambda x: -duplicados[x]):
        freq = duplicados[placeholder]
        print(f"  '{placeholder}' → {freq} vezes")
        # Mostrar contextos
        contextos = [item for item in placeholders_por_posicao if item['placeholder'] == placeholder]
        for ctx in contextos:
            print(f"    - {ctx['tipo']}: {ctx['contexto'][:80]}")
else:
    print("\n✅ Nenhum placeholder duplicado encontrado!")

print("\n" + "=" * 90)
