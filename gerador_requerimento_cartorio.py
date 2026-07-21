import json
import logging
import os
import re
from datetime import datetime
from typing import List, Dict, Any
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

logger = logging.getLogger(__name__)


# ============================================================
# NOMES FEMININOS COMUNS (para detecção de gênero)
# ============================================================
NOMES_FEMININOS = {
    "maria", "ana", "carla", "fernanda", "juliana", "camila",
    "patricia", "andrea", "adriana", "claudia", "denise",
    "elisabete", "eliane", "fabiana", "gabriela", "helena",
    "inez", "jaqueline", "katia", "laura", "lucia", "margarete",
    "marcia", "marta", "monica", "nair", "olivia", "paula",
    "quiteria", "rosa", "silvana", "solange", "teresa",
    "vania", "walquiria", "yasmin", "zuleica", "beatriz",
    "carolina", "daniela", "edite", "fatima", "gisele", "heloisa",
    "isabela", "janice", "lilian", "margaret", "neusa", "odete",
    "priscila", "raquel", "sandra", "tatiane", "ursula", "vitoria",
    "wilma", "yolanda", "zilda", "aline", "bianca", "cristina",
    "diana", "evelin", "flavia", "graziela", "hellen", "ingrid",
    "jessica", "karen", "leticia", "mirela", "natasha", "pamela",
    "renata", "sabrina", "vivian", "angela", "barbara", "cecilia",
    "dolores", "estela", "francisca", "gloria", "hilda", "irma",
    "josefina", "kelly", "lorena", "marina", "paulina", "rita",
    "sueli", "tania", "vera", "wanda", "xenia", "zoraide",
    "amanda", "brenda", "clarice", "daisy", "elenice", "fani",
    "gina", "hebe", "iracema", "joyce", "kika", "melissa",
    "natalia", "rafaela", "silvia", "taina", "vanessa",
    "agostinha", "aparecida", "bete", "creuza", "dalva", "edna",
    "filomena", "gracinda", "ilza", "jandira", "kely", "luana",
    "miriam", "nara", "roseli", "simone", "vanda", "waleska",
    "yanira", "zelia", "ana clara", "vera lucia"
}


def detectar_genero_por_nome(nome: str) -> str:
    """Detecta o gênero de uma pessoa com base no primeiro nome."""
    if not nome or nome.strip().lower() in ("", "xxxxx", "x", "n/a", "-"):
        return "neutro"
    primeiro_nome = nome.strip().lower().split()[0]
    if primeiro_nome in NOMES_FEMININOS:
        return "feminino"
    return "masculino"


def ajustar_profissao_por_genero(profissao: str, genero: str) -> str:
    """Ajusta a profissão conforme o gênero da pessoa."""
    if not profissao:
        return profissao
    if genero == "feminino":
        conjugacoes = {
            "lavrador": "lavradora",
            "agricultor": "agricultora",
            "pecuarista": "pecuarista",
            "engenheiro": "engenheira",
            "tecnico": "técnica",
            "técnico": "técnica",
            "professor": "professora",
            "enfermeiro": "enfermeira",
            "empresario": "empresária",
            "aposentado": "aposentada",
            "desempregado": "desempregada",
            "trabalhador": "trabalhadora",
            "proprietario": "proprietária",
            "proprietário": "proprietária",
            "doutor": "doutora",
            "advogado": "advogada",
            "contador": "contadora",
            "autonomo": "autônoma",
            "autônomo": "autônoma",
        }
        profissao_lower = profissao.lower()
        for masc, fem in conjugacoes.items():
            if masc in profissao_lower:
                # Substituição única e cuidadosa
                profissao = re.sub(re.escape(masc), fem, profissao, flags=re.IGNORECASE, count=1)
                return profissao
    return profissao


def remover_duplicacoes(texto: str) -> str:
    """Remove palavras duplicadas consecutivas no texto."""
    if not texto:
        return texto
    palavras = texto.split()
    resultado = [palavras[0]]
    for i in range(1, len(palavras)):
        if palavras[i].lower() != palavras[i-1].lower():
            resultado.append(palavras[i])
    return " ".join(resultado)


class GeradorRequerimentoCartorio:
    def __init__(self, model_name: str):
        self.model_name = model_name
        self.model = genai.GenerativeModel(model_name)

    def extrair_dados_documentos(self, imagens: List[Any]) -> Dict[str, Any]:
        """
        Envia as imagens dos documentos para o Gemini e extrai os dados estruturados.
        """
        prompt = """
        Analise as imagens dos documentos fornecidos e extraia as informações para um requerimento de cartório.
        REGRAS IMPORTANTES:
        1. Formate RGs com pontos: 706786 -> 706.786 ou 1706786 -> 1.706.786.
        2. Identifique o sexo do requerente para ajustar 'lavrador/lavradora' ou 'agricultor/agricultora'.
           Se for feminino, use 'lavradora' e 'agricultora'. Se masculino, use 'lavrador' e 'agricultor'.
        3. Se não encontrar o cônjuge (esposa), preencha TODOS os campos do requerente_2 com 'XXXXXX'.
        4. Identifique a TRT (começa com BR e tem 11 números).
        5. Identifique a Comarca, Município e Matrícula.
        6. Extraia a área total retificada (encontrada na planta INCRA).
        7. IMPORTANTE: Se encontrar palavras duplicadas como 'comunhão comunhão', corrija para apenas 'comunhão'.
        8. IMPORTANTE: Remova qualquer 'X' usado como marcador de gênero (ex: 'Xlavradora' deve ser 'lavrador' ou 'lavradora' dependendo do gênero).
        9. Se encontrar 2 pessoas (casal), classifique como requerente_1 (proprietário) e requerente_2 (cônjuge/esposa).
        10. Se encontrar apenas 1 pessoa, classifique como requerente_1 e preencha requerente_2 com XXXXXX.
        11. Para o regime de bens, remova duplicações: 'comunhão comunhão de bens' deve virar 'comunhão de bens'.
        
        Retorne estritamente em JSON:
        {
            "comarca": "NOME DA COMARCA EM MAIUSCULO",
            "municipio_cliente": "Cidade do Cliente",
            "requerente_1": {
                "nome": "Nome Completo",
                "profissao": "Lavrador ou Lavradora (ajustado ao gênero)",
                "rg": "0.000.000",
                "orgao": "SSP/ES",
                "cpf": "000.000.000-00",
                "endereco_corrego": "Nome do Córrego (Apenas o nome)"
            },
            "requerente_2": {
                "nome": "Nome Completo da Esposa ou XXXXXX",
                "profissao": "Lavradora ou XXXXXX (ajustado ao gênero)",
                "rg": "0.000.000 ou XXXXXX",
                "orgao": "SSP/ES ou XXXXXX",
                "cpf": "000.000.000-00 ou XXXXXX",
                "regime_bens": "Comunhão Parcial de Bens ou XXXXXX (sem duplicações)"
            },
            "imovel": {
                "nome": "Nome do Sítio (Apenas o nome)",
                "area_registrada": "0,0000",
                "municipio_imovel": "Vila Valério",
                "comarca_imovel": "SÃO GABRIEL DA PALHA",
                "matricula": "0.000",
                "area_encontrada": "0,0000",
                "codigo_incra": "000.000.000.000-0",
                "trt_numero": "BR00000000000",
                "area_total_retificada": "0,0000",
                "cfta_tecnico": "1119851971-1 ou encontrado",
                "codigo_credenciamento": "G1D ou encontrado"
            }
        }
        """
        
        try:
            # Construir conteúdo no formato correto para o Gemini
            logger.info(f"Iniciando extração com {len(imagens)} imagem(s) usando modelo: {self.model_name}")
            conteudo = [prompt]
            for img in imagens:
                conteudo.append(img)
            
            logger.info(f"Enviando requisição para Gemini com {len(conteudo)} elementos (1 prompt + {len(imagens)} imagens)")
            response = self.model.generate_content(conteudo, stream=False)
            
            logger.info(f"Resposta recebida do Gemini. Tamanho: {len(response.text)} caracteres")
            text_response = response.text
            logger.debug(f"Texto bruto da resposta: {text_response[:500]}...")
            if "```json" in text_response:
                text_response = text_response.split("```json")[1].split("```")[0]
            elif "```" in text_response:
                text_response = text_response.split("```")[1].split("```")[0]
            
            logger.info(f"Parseando JSON da resposta...")
            dados = json.loads(text_response.strip())
            logger.info(f"JSON parseado com sucesso. Chaves: {list(dados.keys())}")
            
            # Pós-processamento: ajustar gênero e profissão
            if dados.get("requerente_2"):
                req2 = dados["requerente_2"]
                nome_req2 = req2.get("nome", "")
                if nome_req2.lower() not in ("", "xxxxx", "x", "n/a", "-"):
                    genero = detectar_genero_por_nome(nome_req2)
                    req2["profissao"] = ajustar_profissao_por_genero(req2.get("profissao", ""), genero)
            
            if dados.get("requerente_1"):
                req1 = dados["requerente_1"]
                genero1 = detectar_genero_por_nome(req1.get("nome", ""))
                req1["profissao"] = ajustar_profissao_por_genero(req1.get("profissao", ""), genero1)
            
            # Remover duplicações do regime de bens
            if dados.get("requerente_2", {}).get("regime_bens"):
                dados["requerente_2"]["regime_bens"] = remover_duplicacoes(dados["requerente_2"]["regime_bens"])
            
            logger.info(f"Extração concluída com sucesso")
            return dados
        except Exception as e:
            logger.error(f"Erro na extração via Gemini: {e}", exc_info=True)
            raise Exception(f"Falha ao extrair dados: {str(e)}")

    def _aplicar_negrito_informacoes(self, doc, dados: Dict[str, Any]):
        """Aplica negrito em todos os dados importantes do documento."""
        req1 = dados.get("requerente_1", {})
        req2 = dados.get("requerente_2", {})
        imovel = dados.get("imovel", {})
        
        # Lista completa de dados a colocar em negrito
        dados_negrito = [
            # Requerente 1
            req1.get('nome', ''),
            req1.get('profissao', ''),
            req1.get('rg', ''),
            req1.get('cpf', ''),
            # Requerente 2
            req2.get('nome', ''),
            req2.get('profissao', ''),
            req2.get('rg', ''),
            req2.get('cpf', ''),
            req2.get('regime_bens', ''),
            req1.get('endereco_corrego', ''),
            # Imóvel
            imovel.get('nome', ''),
            imovel.get('area_registrada', ''),
            imovel.get('area_encontrada', ''),
            imovel.get('area_total_retificada', ''),
            imovel.get('matricula', ''),
            imovel.get('municipio_imovel', ''),
            imovel.get('comarca_imovel', ''),
            imovel.get('codigo_incra', ''),
            imovel.get('trt_numero', ''),
        ]
        
        # Remover strings vazias
        dados_negrito = [str(d).strip() for d in dados_negrito if d and str(d).strip() not in ('XXXXXX', 'XXXXX', 'XXX', '')]
        
        # Aplicar negrito em parágrafos
        for para in doc.paragraphs:
            for run in para.runs:
                # Verificar se o texto do run é uma informação importante
                for info in dados_negrito:
                    if info and info.lower() in run.text.lower():
                        # Aplicar negrito apenas neste run
                        run.font.bold = True
                        break
        
        # Aplicar negrito em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for info in dados_negrito:
                                if info and info.lower() in run.text.lower():
                                    run.font.bold = True
                                    break



    def _ajustar_fonte_arial(self, doc):
        """Ajusta a fonte de todo o documento para Arial."""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)

    def _substituir_em_run(self, text_element, substituicoes: Dict[str, str]):
        """
        Substitui placeholders em um elemento de texto (parágrafo ou cell).
        Trabalha no nível dos RUNS para preservar formatação.
        """
        # Junta todo o texto do elemento
        texto_completo = ""
        for run in text_element.runs:
            texto_completo += run.text
        
        # Verifica se há algum placeholder para substituir
        tem_placeholder = False
        for key in substituicoes:
            if key in texto_completo:
                tem_placeholder = True
                break
        
        if not tem_placeholder:
            return
        
        # Aplica todas as substituições no texto completo
        for key, val in substituicoes.items():
            texto_completo = texto_completo.replace(key, str(val))
        
        # Reescreve o texto: primeiro run recebe tudo, demais ficam vazios
        if text_element.runs:
            text_element.runs[0].text = texto_completo
            for run in text_element.runs[1:]:
                run.text = ""

    def gerar_documento(self, dados: Dict[str, Any], template_name: str) -> bytes:
        try:
            # Busca do template
            base_path = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(base_path, template_name)
            if not os.path.exists(template_path):
                template_path = template_name
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Modelo {template_name} não encontrado.")

            doc = Document(template_path)
            
            # Data de hoje formatada
            hoje = datetime.now()
            meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
                     "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
            data_formatada = f"{hoje.day} de {meses[hoje.month-1]} de {hoje.year}"

            # Extrair dados com valores padrão seguros
            req1 = dados.get("requerente_1", {})
            req2 = dados.get("requerente_2", {})
            imovel = dados.get("imovel", {})

            # Determinar estado civil
            nome_req2 = req2.get("nome", "").strip().lower()
            tem_conjuge = nome_req2 not in ("", "xxxxx", "x", "n/a", "-")
            estado_civil_req1 = "casado" if tem_conjuge else "solteiro"

            logger.info(f"Gerando documento com estado civil: {estado_civil_req1}, tem cônjuge: {tem_conjuge}")

            # ============================================================
            # MAPEAMENTO DE PLACEHOLDERS - Template com ( )
            # Usando os placeholders exatos do template Word
            # ============================================================
            # Preparar profissão com vírgula
            profissao_req1 = req1.get('profissao', 'XXXXX')
            profissao_req2 = req2.get('profissao', 'XXXXX')
            
            # Preparar regime de bens (remover duplicações)
            regime_bens = remover_duplicacoes(req2.get('regime_bens', 'XXXXXX'))
            
            # Preparar endereço (remover duplicações)
            endereco_corrego = remover_duplicacoes(req1.get('endereco_corrego', 'XXXXX'))
            
            # Preparar nome do imóvel (remover duplicações)
            nome_imovel = remover_duplicacoes(imovel.get('nome', 'XXXXX'))
            
            # Áreas das glebas
            area_gleba_1 = imovel.get('area_registrada', 'XXXXXX')  # Gleba 1 = área registrada
            area_gleba_2 = imovel.get('area_encontrada', 'XXXXX')   # Gleba 2 = área encontrada
            
            substituicoes = {
                # CABEÇALHO / DESTINATÁRIO
                "(XXXXX)": imovel.get('comarca_imovel', 'XXXXX'),
                
                # REQUERENTE 1 - Dados principais
                "(NOME_REQUERENTE_1)": req1.get('nome', 'XXXXXX'),
                "(ESTADO_CIVIL_REQUERENTE_1)": f"{estado_civil_req1},",  # Adicionar vírgula
                "(PROFISSAO_REQUERENTE_1)": profissao_req1,
                "(RG_REQUERENTE_1)": req1.get('rg', 'XXXX'),
                "(ORGAO_RG_REQUERENTE_1)": req1.get('orgao', 'SSP/ES'),
                "(CPF_REQUERENTE_1)": req1.get('cpf', 'XXXXXXX'),
                "(ENDERECO_CORREGO_REQUERENTE_1)": endereco_corrego,
                
                # REQUERENTE 2 - Dados principais
                "(NOME_REQUERENTE_2)": req2.get('nome', 'XXXXXX'),
                "(PROFISSAO_REQUERENTE_2)": profissao_req2,
                "(RG_REQUERENTE_2)": req2.get('rg', 'XXXXX'),
                "(ORGAO_RG_REQUERENTE_2)": req2.get('orgao', 'SSP/ES'),
                "(CPF_REQUERENTE_2)": req2.get('cpf', 'XXXXXX'),
                "(REGIME_BENS)": regime_bens,
                
                # LOCALIZAÇÃO
                "(CIDADE DO REQUERENTE -ES)": f"{dados.get('municipio_cliente', 'XXXXXX')}-ES",
                
                # IMÓVEL - Identificação
                "(NOME_IMOVEL)": nome_imovel,
                "(AREA DO IMOVEL há)": imovel.get('area_registrada', 'XXXXXX'),
                "(CIDADE ONDE O IMOVEL ESTÁ REGISTRADO – ES)": f"{imovel.get('municipio_imovel', 'XXXX')}-ES",
                "(COMARCA ONDE O IMOVEL ESTÁ REGISTRADO – ES)": imovel.get('comarca_imovel', 'XXXXXXX'),
                "(MATRICULA DO IMOVEL)": imovel.get('matricula', 'XXXXXX'),
                
                # IMÓVEL - Valores
                "(VALOR DO IMOVEL EM NUMEROS)": imovel.get('valor_fiscal', 'XX0.000,00'),
                "(VALOR DO IMOVEL POR ESCRITO mil reais)": "XXXXXX",  # Não extraído
                
                # IMÓVEL - Áreas
                "(ÁREA ENCONTRADA NA PLANTA DO IMOVEL há)": imovel.get('area_encontrada', 'XXXXX'),
                "(XXXXXX há)": area_gleba_1,  # Gleba 1 = área registrada
                "(XXXXX há)": area_gleba_2,   # Gleba 2 = área encontrada
                "(ÁREA DO IMOVEL há)": imovel.get('area_total_retificada', 'XXXXXXX'),  # Retificação
                
                # IMÓVEL - Certificações
                "(CÓDIGO DO IMÓVEL RURAL DISPONIVEL NO CCIR)": imovel.get('codigo_incra', 'XXX.XXX.XXX.XXX-X'),
                "(BRXXXXXXX(TRT ENCONTRADA EM TRT)": imovel.get('trt_numero', 'BRXXXXXXX').rstrip(')'),  # Remover parêntese extra
                
                # DATA
                "(XX de XXX de XXXX)": data_formatada,
            }

            logger.info(f"Aplicando {len(substituicoes)} substituições no documento...")

            # Aplicando substituições nos parágrafos (nível de run)
            for p in doc.paragraphs:
                self._substituir_em_run(p, substituicoes)
            
            # Aplicando substituições nas tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._substituir_em_run(p, substituicoes)
            
            # Ajuste final de fonte
            self._ajustar_fonte_arial(doc)
            
            # Aplicar negrito nas informações importantes
            logger.info(f"Aplicando negrito nas informações importantes...")
            self._aplicar_negrito_informacoes(doc, dados)

            logger.info(f"Salvando documento em buffer...")
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Retornar bytes (não BytesIO)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Erro ao gerar documento: {e}", exc_info=True)
            raise Exception(f"Falha ao gerar documento: {str(e)}")
