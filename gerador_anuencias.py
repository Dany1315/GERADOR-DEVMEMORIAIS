#"""
#MÓDULO INDEPENDENTE: GERADOR E ANALISADOR DE ANUÊNCIAS VIA GEMINI API
#Este arquivo roda de forma isolada e não afeta o código do Gerador de Memoriais.
#"""

import io
import re
import logging
from datetime import datetime
from typing import Dict, List, Any
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

logger = logging.getLogger(__name__)

class GeradorAnuenciaWord:
    def __init__(self, dados_empresa: Dict[str, str], dados_tecnico: Dict[str, str]):
        """
        Inicializa o gerador com os dados padrão da empresa e do responsável técnico.
        """
        self.dados_empresa = dados_empresa
        self.dados_tecnico = dados_tecnico
        
        # Configuração da API do Gemini (Mantido por compatibilidade de estrutura)
        self.api_key = st.secrets.get("GEMINI_API_KEY")
        if self.api_key:
            genai.configure(api_key=self.api_key)

    def _limpar_e_converter_numerico(self, valor: Any) -> float:
        """
        Limpa caracteres indesejados mantendo apenas números, pontos e vírgulas,
        garantindo a conversão para float sem erros de compilação.
        """
        try:
            string_limpa = str(valor).strip()
            string_limpa = re.sub(r"[^\d,.]", "", string_limpa)
            if not string_limpa:
                return 0.0
            
            # Se possui os dois separadores (ex: 7,884,471.06), remove as vírgulas de milhar
            if "," in string_limpa and "." in string_limpa:
                string_limpa = string_limpa.replace(",", "")
            # Caso tenha vindo apenas com vírgula como decimal (ex: 15,68)
            elif "," in string_limpa and "." not in string_limpa:
                string_limpa = string_limpa.replace(",", ".")
                
            return float(string_limpa)
        except Exception:
            return 0.0

    def _formatar_para_padrao_modelo(self, valor_numerico: float) -> str:
        """
        Formata o número usando ponto para decimais e vírgula para milhar (Ex: 7,884,471.06).
        """
        # Formata com vírgulas como separador de milhar americano (1,234,567.89)
        return f"{valor_numerico:,.2f}"

    def gerar_documento(self, dados_anuencia: Dict[str, Any]) -> io.BytesIO:
        """
        Gera o arquivo Word (.docx) baseado estritamente no modelo físico de anuência enviado.
        """
        confrontante = dados_anuencia["confrontante"]
        proprietario = dados_anuencia["proprietario"]
        local = dados_anuencia.get("local", "Vila Valério")
        segmentos = dados_anuencia["segmentos"]

        doc = Document()

        # Configuração de Margens Padrão (2.54 cm / 1 polegada)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Configuração de Estilo e Fonte (Calibri 11)
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # =========================================================================
        # CABEÇALHO CORPORATIVO (Topo da Folha, Centralizado)
        # =========================================================================
        p_head1 = doc.add_paragraph()
        p_head1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head1.paragraph_format.space_before = Pt(0)
        p_head1.paragraph_format.space_after = Pt(1)
        run_h1 = p_head1.add_run("TopoGeo")
        run_h1.bold = True
        run_h1.font.size = Pt(12)
        
        p_head2 = doc.add_paragraph()
        p_head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head2.paragraph_format.space_before = Pt(0)
        p_head2.paragraph_format.space_after = Pt(1)
        run_h2 = p_head2.add_run("Topografia   Consultoria LTDA")
        run_h2.bold = True
        run_h2.font.size = Pt(10)
        
        p_head3 = doc.add_paragraph()
        p_head3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head3.paragraph_format.space_before = Pt(0)
        p_head3.paragraph_format.space_after = Pt(1)
        run_h3 = p_head3.add_run("Rua Natalino Cossi, No 114, sala 2 - Vila Valério, CEP 29785-000 Fone 27 99837-1164")
        run_h3.font.size = Pt(9)
        run_h3.font.color.rgb = RGBColor(100, 100, 100) # Cor cinza elegante
        
        p_head4 = doc.add_paragraph()
        p_head4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_head4.paragraph_format.space_before = Pt(0)
        p_head4.paragraph_format.space_after = Pt(24) # Espaçamento maior antes do título
        run_h4 = p_head4.add_run("topogeo2014@gmail.com")
        run_h4.font.size = Pt(9)
        run_h4.font.color.rgb = RGBColor(100, 100, 100)

        # 1. TÍTULO: Negrito, Centralizado e em Letras Maiúsculas
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after = Pt(18)
        run_titulo = p_titulo.add_run("DECLARAÇÃO DE RECONHECIMENTO DE LIMITES")
        run_titulo.bold = True
        run_titulo.font.size = Pt(11)

        # 2. TEXTO DE ABERTURA: Alinhamento Justificado
        texto_abertura = (
            f"Eu, {confrontante.title()}, proprietário do imóvel confrontante, e eu, "
            f"{proprietario.title()}, proprietário do imóvel urbano, declaramos não "
            f"existir nenhuma disputa ou discordância sobre os limites comuns existentes entre os citados imóveis."
        )
        p_abertura = doc.add_paragraph(texto_abertura)
        p_abertura.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_abertura.paragraph_format.line_spacing = 1.15
        p_abertura.paragraph_format.space_after = Pt(12)

        # 3. DESCRIÇÃO DO TRECHO (Vai direto do Título para a Tabela, sem parágrafo de IA)
        p_desc_tit = doc.add_paragraph()
        p_desc_tit.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_desc_tit.paragraph_format.space_after = Pt(6)
        p_desc_tit.add_run("Descrição do trecho de confrontação:").bold = True

        # 4. TABELA TÉCNICA: Mantém linhas de grade visíveis ('Table Grid') e padrão americano
        tabela = doc.add_table(rows=1, cols=7)
        tabela.style = 'Table Grid'
        tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        hdr_cells = tabela.rows[0].cells
        headers = ["De", "Para", "Azimute", "Distância (m)", "E(X)", "N(Y)", "Altitude"]
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.bold = True
            run.font.size = Pt(10)

        total_distancia = 0.0
        for s in segmentos:
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(s['de'])
            row_cells[1].text = str(s['para'])
            row_cells[2].text = str(s['azimute'])
            
            # Tratamento da distância e acumulação
            dist_val = self._limpar_e_converter_numerico(s['distancia'])
            total_distancia += dist_val
            row_cells[3].text = self._formatar_para_padrao_modelo(dist_val)
            
            # Formatação de Coordenadas sem sufixos de texto (ex: "m")
            val_ex = self._limpar_e_converter_numerico(s.get('e_x', '0.00'))
            val_ny = self._limpar_e_converter_numerico(s.get('n_y', '0.00'))
            
            row_cells[4].text = self._formatar_para_padrao_modelo(val_ex)
            row_cells[5].text = self._formatar_para_padrao_modelo(val_ny)
            row_cells[6].text = "0,00"
            
            # Centralizar células de dados da tabela
            for cell in row_cells:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.size = Pt(9.5)

        # Linha de Totais da Tabela (Sem espaços e alinhada)
        row_totais = tabela.add_row().cells
        row_totais[0].text = "Total"
        row_totais[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_totais[0].paragraphs[0].runs[0].font.bold = True
        row_totais[0].paragraphs[0].runs[0].font.size = Pt(9.5)

        row_totais[3].text = self._formatar_para_padrao_modelo(total_distancia)
        row_totais[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_totais[3].paragraphs[0].runs[0].font.bold = True
        row_totais[3].paragraphs[0].runs[0].font.size = Pt(9.5)
        
        # Células vazias limpas
        for idx in [1, 2, 4, 5, 6]:
            row_totais[idx].text = ""
            row_totais[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Espaço pós tabela
        p_espaco = doc.add_paragraph("")
        p_espaco.paragraph_format.space_before = Pt(6)

        # 5. PARÁGRAFO DE RESPONSABILIDADE TÉCNICA: Justificado e com colagem sem espaçamento pós-ponto
        rg_rt = self.dados_tecnico.get('rg', '1.936.653')
        codigo_incra = self.dados_tecnico.get('codigo_incra', 'G1D')
        
        texto_tecnico = (
            f"Declaramos ainda que o profissional {self.dados_tecnico.get('nome')} "
            f"(RG nº {rg_rt} e CPF nº {self.dados_tecnico.get('cpf', '111.985.197-11')}), Resp. "
            f"Técnico (CFTA {self.dados_tecnico.get('cfta')}), credenciado pelo INCRA sob o cod. {codigo_incra}, com a emissão da TRT nº "
            f"{self.dados_tecnico.get('trt')}, nos indicou as demarcações do limite entre as nossas propriedades, tanto no campo como "
            f"nas suas apresentações gráficas.Concordamos com essa demarcação, expressa na planta e no memorial descritivo, "
            f"ambos em anexo, e reconhecemos esta descrição como o limite legal entre nossas propriedades."
        )
        p_tecnico = doc.add_paragraph(texto_tecnico)
        p_tecnico.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_tecnico.paragraph_format.line_spacing = 1.15
        p_tecnico.paragraph_format.space_after = Pt(20)

        # 6. ENCERRAMENTO COM DATA
        data_atual = datetime.now().strftime('%d de %m de %Y')
        p_data = doc.add_paragraph(f"{local}, {data_atual}")
        p_data.paragraph_format.space_after = Pt(28)

        # 7. CAMPOS DE ASSINATURA LADO A LADO: Tabela Invisível de 2 colunas
        tab_assinatura = doc.add_table(rows=2, cols=2)
        tab_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Esconder bordas para a seção de assinatura ficar perfeitamente limpa
        tblPr = tab_assinatura._tbl.tblPr
        tblBorders = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="none"/>'
            r'  <w:left w:val="none"/>'
            r'  <w:bottom w:val="none"/>'
            r'  <w:right w:val="none"/>'
            r'  <w:insideH w:val="none"/>'
            r'  <w:insideV w:val="none"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(tblBorders)

        # Linhas de assinatura (Underlines encurtados de 50 para 38 caracteres)
        celulas_l1 = tab_assinatura.rows[0].cells
        celulas_l1[0].text = "______________________________________"
        celulas_l1[1].text = "______________________________________"
        celulas_l1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celulas_l1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nomes e cargos em parágrafos separados (evita fusão de linhas)
        celulas_l2 = tab_assinatura.rows[1].cells
        
        # Coluna 1: Confrontante
        c1 = celulas_l2[0]
        c1.text = "" # Limpa inicialização
        p1_c1 = c1.paragraphs[0]
        p1_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1_c1.paragraph_format.space_before = Pt(4)
        run_name_c1 = p1_c1.add_run(confrontante.title())
        run_name_c1.font.size = Pt(10)
        
        p2_c1 = c1.add_paragraph()
        p2_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2_c1.paragraph_format.space_before = Pt(1)
        run_role_c1 = p2_c1.add_run("Proprietário do Imóvel Confrontante")
        run_role_c1.font.size = Pt(10)

        # Coluna 2: Proprietário
        c2 = celulas_l2[1]
        c2.text = "" # Limpa inicialização
        p1_c2 = c2.paragraphs[0]
        p1_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1_c2.paragraph_format.space_before = Pt(4)
        run_name_c2 = p1_c2.add_run(proprietario.title())
        run_name_c2.font.size = Pt(10)
        
        p2_c2 = c2.add_paragraph()
        p2_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2_c2.paragraph_format.space_before = Pt(1)
        run_role_c2 = p2_c2.add_run("Proprietário do Imóvel")
        run_role_c2.font.size = Pt(10)

        # Espaçamento para o RT
        p_espaco_final = doc.add_paragraph("")
        p_espaco_final.paragraph_format.space_before = Pt(16)

        # 8. ASSINATURA DO RESPONSÁVEL TÉCNICO: Centralizado e estruturado em 4 parágrafos separados
        # Linha de assinatura do RT (Underlines encurtados de 30 para 24 caracteres)
        p_rt_l1 = doc.add_paragraph()
        p_rt_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rt_l1.paragraph_format.space_before = Pt(12)
        p_rt_l1.paragraph_format.space_after = Pt(0)
        run_linha = p_rt_l1.add_run("________________________")
        run_linha.bold = True
        
        p_rt_l2 = doc.add_paragraph()
        p_rt_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rt_l2.paragraph_format.space_before = Pt(4)
        p_rt_l2.paragraph_format.space_after = Pt(0)
        run_nome_tec = p_rt_l2.add_run(f"{self.dados_tecnico.get('nome')}")
        run_nome_tec.bold = True
        
        p_rt_l3 = doc.add_paragraph()
        p_rt_l3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rt_l3.paragraph_format.space_before = Pt(1)
        p_rt_l3.paragraph_format.space_after = Pt(0)
        run_cargo = p_rt_l3.add_run("Resp. Técnico")
        run_cargo.font.size = Pt(10)
        
        p_rt_l4 = doc.add_paragraph()
        p_rt_l4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_rt_l4.paragraph_format.space_before = Pt(1)
        p_rt_l4.paragraph_format.space_after = Pt(0)
        run_cfta = p_rt_l4.add_run(f"CFTA: {self.dados_tecnico.get('cfta')}")
        run_cfta.font.size = Pt(10)

        # Salva o arquivo final para download
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
