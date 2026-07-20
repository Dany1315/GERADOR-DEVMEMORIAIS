#!/usr/bin/env python3
"""
Script para criar um novo template de requerimento com placeholders únicos.
Este script lê o template original e substitui todos os placeholders duplicados
por placeholders únicos e descritivos.
"""

from docx import Document
import re
from collections import defaultdict

# Mapeamento de placeholders antigos (duplicados) para novos (únicos)
MAPEAMENTO_PLACEHOLDERS = {
    # Nome proprietário (1ª ocorrência de (XXXXXX))
    "proprietario_nome": {
        "antigo": "(XXXXXX)",
        "novo": "{{NOME_PROPRIETARIO}}",
        "contexto": "proprietario"
    },
    
    # RG proprietário (1ª ocorrência de (XXXXX))
    "proprietario_rg": {
        "antigo": "(XXXXX)",
        "novo": "{{RG_PROPRIETARIO}}",
        "contexto": "rg"
    },
    
    # CPF proprietário (1ª ocorrência de XXXXXXX sem parênteses)
    "proprietario_cpf": {
        "antigo": "XXXXXXX",
        "novo": "{{CPF_PROPRIETARIO}}",
        "contexto": "cpf"
    },
    
    # Nome esposa (2ª ocorrência de (XXXXXX))
    "esposa_nome": {
        "antigo": "(XXXXXX)",
        "novo": "{{NOME_ESPOSA}}",
        "contexto": "esposa"
    },
    
    # RG esposa (2ª ocorrência de (XXXXX))
    "esposa_rg": {
        "antigo": "(XXXXX)",
        "novo": "{{RG_ESPOSA}}",
        "contexto": "esposa"
    },
    
    # CPF esposa (2ª ocorrência de (XXXXXX) após esposa)
    "esposa_cpf": {
        "antigo": "(XXXXXX)",
        "novo": "{{CPF_ESPOSA}}",
        "contexto": "esposa"
    },
    
    # Regime de bens
    "regime_bens": {
        "antigo": "(XXXXXX)",
        "novo": "{{REGIME_BENS}}",
        "contexto": "comunhão"
    },
    
    # Córrego/Endereço
    "endereco_corrego": {
        "antigo": "(XXXXX)",
        "novo": "{{ENDERECO_CORREGO}}",
        "contexto": "córrego"
    },
    
    # Nome sítio
    "nome_sitio": {
        "antigo": "(XXXXX)",
        "novo": "{{NOME_SITIO}}",
        "contexto": "sítio"
    },
    
    # Área registrada
    "area_registrada": {
        "antigo": "(XXXXXX há)",
        "novo": "{{AREA_REGISTRADA}} há",
        "contexto": "registrada"
    },
    
    # Município imóvel
    "municipio_imovel": {
        "antigo": "(XXXX – ES)",
        "novo": "{{MUNICIPIO_IMOVEL}} – ES",
        "contexto": "município"
    },
    
    # Comarca imóvel
    "comarca_imovel": {
        "antigo": "(XXXXXXX – ES)",
        "novo": "{{COMARCA_IMOVEL}} – ES",
        "contexto": "comarca"
    },
    
    # Matrícula
    "matricula": {
        "antigo": "(XXXXXX)",
        "novo": "{{MATRICULA}}",
        "contexto": "matrícula"
    },
    
    # Área encontrada
    "area_encontrada": {
        "antigo": "(XXXXX há)",
        "novo": "{{AREA_ENCONTRADA}} há",
        "contexto": "encontrada"
    },
    
    # Código INCRA
    "codigo_incra": {
        "antigo": "(XXX.XXX.XXX.XXX-X)",
        "novo": "{{CODIGO_INCRA}}",
        "contexto": "incra"
    },
    
    # TRT número
    "trt_numero": {
        "antigo": "(BRXXXXXXX)",
        "novo": "{{TRT_NUMERO}}",
        "contexto": "trt"
    },
    
    # Área total retificada
    "area_total_retificada": {
        "antigo": "(XXXXXXX há)",
        "novo": "{{AREA_TOTAL_RETIFICADA}} há",
        "contexto": "total retificada"
    },
    
    # Data
    "data": {
        "antigo": "(XX de XXX de XXXX)",
        "novo": "{{DATA_FORMATADA}}",
        "contexto": "data"
    },
    
    # Assinaturas
    "assinatura_proprietario": {
        "antigo": "(XXXXXXXXXXXXXXXX)",
        "novo": "{{ASSINATURA_PROPRIETARIO}}",
        "contexto": "assinatura"
    },
    
    "assinatura_esposa": {
        "antigo": "(XXXXXXXXXXXXXXXXXX)",
        "novo": "{{ASSINATURA_ESPOSA}}",
        "contexto": "esposa"
    },
}


def substituir_placeholders_no_template(template_path, output_path):
    """
    Lê o template original e substitui placeholders duplicados por únicos.
    """
    print("=" * 80)
    print("CRIANDO TEMPLATE COM PLACEHOLDERS ÚNICOS")
    print("=" * 80)
    
    doc = Document(template_path)
    
    # Contador de substituições
    substituicoes_realizadas = defaultdict(int)
    
    # Processar parágrafos
    print("\n📄 Processando parágrafos...")
    for para_idx, p in enumerate(doc.paragraphs):
        texto_original = p.text
        texto_novo = texto_original
        
        # Substituir cada placeholder
        for campo, info in MAPEAMENTO_PLACEHOLDERS.items():
            antigo = info['antigo']
            novo = info['novo']
            contexto = info['contexto']
            
            # Verificar se o placeholder existe e se o contexto corresponde
            if antigo in texto_novo and contexto.lower() in texto_novo.lower():
                # Substituir apenas a primeira ocorrência
                texto_novo = texto_novo.replace(antigo, novo, 1)
                substituicoes_realizadas[campo] += 1
        
        # Atualizar parágrafo se houve mudanças
        if texto_novo != texto_original:
            # Limpar runs existentes
            for run in p.runs:
                run.text = ""
            
            # Adicionar novo texto
            if p.runs:
                p.runs[0].text = texto_novo
            else:
                p.add_run(texto_novo)
            
            print(f"  ✓ Parágrafo {para_idx}: {len(texto_original)} → {len(texto_novo)} caracteres")
    
    # Processar tabelas
    print("\n📊 Processando tabelas...")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, p in enumerate(cell.paragraphs):
                    texto_original = p.text
                    texto_novo = texto_original
                    
                    # Substituir cada placeholder
                    for campo, info in MAPEAMENTO_PLACEHOLDERS.items():
                        antigo = info['antigo']
                        novo = info['novo']
                        contexto = info['contexto']
                        
                        if antigo in texto_novo and contexto.lower() in texto_novo.lower():
                            texto_novo = texto_novo.replace(antigo, novo, 1)
                            substituicoes_realizadas[campo] += 1
                    
                    # Atualizar parágrafo se houve mudanças
                    if texto_novo != texto_original:
                        for run in p.runs:
                            run.text = ""
                        
                        if p.runs:
                            p.runs[0].text = texto_novo
                        else:
                            p.add_run(texto_novo)
                        
                        print(f"  ✓ Tabela {table_idx}, Linha {row_idx}, Coluna {cell_idx}: Atualizado")
    
    # Salvar novo template
    doc.save(output_path)
    
    # Relatório
    print("\n" + "=" * 80)
    print("RESUMO DE SUBSTITUIÇÕES")
    print("=" * 80)
    
    for campo in sorted(substituicoes_realizadas.keys()):
        freq = substituicoes_realizadas[campo]
        info = MAPEAMENTO_PLACEHOLDERS[campo]
        print(f"✓ {campo:30} {info['antigo']:20} → {info['novo']:30} ({freq}x)")
    
    print(f"\nTotal de substituições: {sum(substituicoes_realizadas.values())}")
    print(f"\n✅ Novo template salvo em: {output_path}")
    print("=" * 80)


if __name__ == "__main__":
    template_original = "/home/ubuntu/upload/GERADOR-DEVMEMORIAIS-main/-REQUERIMENTODECARTORIO.docx"
    template_novo = "/home/ubuntu/upload/GERADOR-DEVMEMORIAIS-main/-REQUERIMENTODECARTORIO_CORRIGIDO.docx"
    
    substituir_placeholders_no_template(template_original, template_novo)
