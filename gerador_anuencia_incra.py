import io
import re
import json
import logging
import zipfile
from datetime import datetime
from typing import Dict, Any, List, Tuple
import streamlit as st
import google.generativeai as genai

try:
    import pypdf
except ImportError:
    pypdf = None

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

logger = logging.getLogger(__name__)


class GeradorAnuenciaIncraWord:
    def __init__(self, dados_empresa: Dict[str, str], dados_tecnico: Dict[str, str]):
        """
        Inicializa o gerador de anuência do INCRA com dados institucionais e do técnico.
        """
        self.dados_empresa = dados_empresa
        self.dados_tecnico = dados_tecnico

        # Configuração da API do Gemini obtida de forma segura dos secrets do Streamlit
        self.api_key = st.secrets.get("GEMINI_API_KEY")
        if self.api_key:
            genai.configure(api_key=self.api_key)
        else:
            st.warning(
                "GEMINI_API_KEY não encontrada em st.secrets — o gerador vai usar "
                "os dados padrão (fallback) em vez de analisar o memorial com IA."
            )

    def _estrutura_padrao(self) -> Dict[str, Any]:
        """
        Estrutura de fallback baseada fielmente nos modelos reais fornecidos.
        """
        return {
            "proprietario_origem": "RODRIGO COLOMBI FROTA",
            "cpf_origem": "092.653.737-76",
            "confrontantes": [
                {
                    "confrontante_imovel": "Sitio Bravin",
                    "confrontante_matricula": "6093",
                    "confrontante_comarca": "São Gabriel da Palha",
                    "confrontante_proprietario": "ANTONIO BRAVIN",
                    "confrontante_cpf": "___.___.___-__",
                    "vertices": [
                        {
                            "codigo": "G1D-M-03281",
                            "longitude": "40°22'10.639\"",
                            "latitude": "19°00'24.525\"",
                            "altitude": "162.10",
                            "vante": "G1D-M-03282",
                            "azimute": "06°49'",
                            "distancia": "640.71",
                            "confrontacao_completa": "CNS: 02.170-9 | Mat. 6093 | Sitio Bravin; Antonio Bravin"
                        }
                    ]
                }
            ]
        }

    def _formatar_coordenada(self, coord_str: str) -> str:
        """
        Garante que as coordenadas de Latitude e Longitude fiquem limpas de sinais negativos,
        com ponto decimal nas casas decimais e formatação de grau (GG°MM'SS.SSS").
        """
        if not coord_str:
            return ""
        
        # Remove sinal de menos (-) e aspas duplas duplicadas
        limpo = coord_str.replace("-", "").replace('""', '"').strip()
        
        # Troca vírgula por ponto para o padrão decimal
        limpo = limpo.replace(",", ".")
        
        # Expressão regular para capturar os componentes GG, MM, SS.SSS
        match = re.search(r"(\d+)[°ºd\s]+(\d+)['\'\s]+([\d\.]+)", limpo)
        if match:
            graus = match.group(1)
            minutos = match.group(2)
            segundos = float(match.group(3))
            return f"{graus}°{minutos}'{segundos:.3f}\""
            
        # Se não capturar pela Regex estruturada, faz substituições básicas de segurança
        limpo = limpo.replace("d", "°").replace("'", "'").replace('"', '"')
        if "°" not in limpo and len(limpo) > 4:
            # Caso venha um formato numérico corrido "402210.639"
            return f"{limpo[:2]}°{limpo[2:4]}'{limpo[4:]}\""
        return limpo

    def _formatar_azimute(self, az_str: str) -> str:
        """
        Garante a formatação rigorosa do azimute (ex: 06°49').
        """
        if not az_str:
            return ""
        limpo = az_str.replace("-", "").strip().replace(",", ".")
        match = re.search(r"(\d+)[°ºd\s]+(\d+)", limpo)
        if match:
            graus = match.group(1)
            minutos = match.group(2)
            return f"{int(graus):02d}°{int(minutos):02d}'"
        return limpo

    def _formatar_numero(self, num_str: Any, casas: int = 2) -> str:
        """
        Converte qualquer valor numérico ou string com vírgula para float com ponto e string de casas fixas.
        """
        if num_str is None:
            return ""
        try:
            val = str(num_str).replace(",", ".").strip()
            # Extrai apenas o número com ponto decimal
            match = re.search(r"[\d\.]+", val)
            if match:
                f_val = float(match.group(0))
                return f"{f_val:.{casas}f}"
            return str(num_str)
        except Exception:
            return str(num_str)

    def _obter_dados_estruturados_com_ia(self, texto_memorial: str, dados_projeto: Dict[str, Any]) -> Dict[str, Any]:
        """
        Usa o Gemini para analisar o memorial e separar TODOS os confrontantes com seus respectivos
        vértices e dados cadastrais, além de extrair o proprietário do imóvel de origem.
        """
        estrutura_padrao = self._estrutura_padrao()

        if not self.api_key:
            logger.warning("Chave de API do Gemini não configurada nos secrets do Streamlit. Usando dados padrão.")
            return estrutura_padrao

        prompt = f"""
        Você é um engenheiro cartógrafo especialista em georreferenciamento do INCRA.
        Sua tarefa é analisar o texto de um memorial descritivo ou relatório de vértices e estruturar as
        informações de TODOS os confrontantes (vizinhos) identificados ao longo da poligonal.

        Além disso, você deve identificar no texto quem é o PROPRIETÁRIO DO IMÓVEL DE ORIGEM (o dono da terra que está sendo medida/georreferenciada) e seu CPF se houver.

        TEXTO DO MEMORIAL DESCRITIVO / RELATÓRIO EXTRAÍDO:
        \"\"\"
        {texto_memorial}
        \"\"\"

        REGRAS DE EXTRAÇÃO:
        1. Identifique o proprietário de origem (o sujeito da medição principal do memorial) e o seu CPF. Se não encontrar o CPF no texto, use o valor padrão "{dados_projeto.get('cpf_proprietario', '092.653.737-76')}".
        2. Identifique cada confrontante distinto (por nome de proprietário, matrícula ou imóvel) ao longo do perímetro.
        3. Agrupe sob cada confrontante APENAS os vértices/segmentos cujo trecho de "vante" faz divisa com ele.
        4. Se o texto for incompreensível, simule dados verossímeis baseados no contexto para manter o JSON estruturado.
        5. Tente encontrar ou estimar o CPF do confrontante se houver menção, caso contrário, deixe em branco para preenchimento manual (ex: "___.___.___-__").

        Responda APENAS com o JSON estruturado abaixo, sem markdown ou textos explicativos:
        {{
            "proprietario_origem": "NOME DO PROPRIETÁRIO DO IMÓVEL DE ORIGEM ENCONTRADO NO MEMORIAL (EM LETRAS MAIÚSCULAS)",
            "cpf_origem": "CPF DO PROPRIETÁRIO DE ORIGEM ENCONTRADO (OU O PADRÃO ENVIADO)",
            "confrontantes": [
                {{
                    "confrontante_imovel": "Nome do Imóvel Confrontante",
                    "confrontante_matricula": "Número da Matrícula",
                    "confrontante_comarca": "Nome da Comarca",
                    "confrontante_proprietario": "Nome do Proprietário Confrontante",
                    "confrontante_cpf": "CPF do Confrontante",
                    "vertices": [
                        {{
                            "codigo": "Código do Vértice",
                            "longitude": "Longitude formatada",
                            "latitude": "Latitude formatada",
                            "altitude": "Altitude com duas casas",
                            "vante": "Código do Vértice de Vante",
                            "azimute": "Azimute formatado",
                            "distancia": "Distância formatada com duas casas",
                            "confrontacao_completa": "Descrição completa da confrontação"
                        }}
                    ]
                }}
            ]
        }}
        """

        try:
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(
                prompt,
                generation_config={"response_mime_type": "application/json"}
            )

            texto_resposta = response.text.strip()
            if texto_resposta.startswith("```json"):
                texto_resposta = texto_resposta.split("```json")[1].split("```")[0].strip()
            elif texto_resposta.startswith("```"):
                texto_resposta = texto_resposta.split("```")[1].split("```")[0].strip()

            dados = json.loads(texto_resposta)

            if "confrontantes" not in dados:
                dados = {"confrontantes": [dados]}

            return dados
        except Exception as e:
            logger.error(f"Erro ao obter dados estruturados do Gemini: {str(e)}")
            st.error(
                f"Falha ao consultar o Gemini — usando dados padrão (fallback). "
                f"Detalhe do erro: {str(e)}"
            )
            return estrutura_padrao

    def _extrair_texto_memorial(self, conteudo_arquivo: bytes, nome_arquivo: str) -> str:
        texto_memorial = ""
        if nome_arquivo.lower().endswith(".pdf"):
            try:
                pdf_file = io.BytesIO(conteudo_arquivo)
                if pypdf:
                    reader = pypdf.PdfReader(pdf_file)
                    paginas_texto = []
                    for pagina in reader.pages:
                        texto_pag = pagina.extract_text()
                        if texto_pag:
                            paginas_texto.append(texto_pag)
                    texto_memorial = "\n".join(paginas_texto)
                else:
                    texto_memorial = conteudo_arquivo.decode("utf-8", errors="ignore")
            except Exception as pdf_err:
                logger.error(f"Falha ao extrair texto do PDF: {pdf_err}")
                texto_memorial = "Erro na leitura estruturada do PDF."
        elif nome_arquivo.lower().endswith(".docx"):
            try:
                doc_temp = Document(io.BytesIO(conteudo_arquivo))
                texto_memorial = "\n".join([p.text for p in doc_temp.paragraphs])
            except Exception as docx_err:
                logger.error(f"Não foi possível abrir o arquivo como .docx: {docx_err}")
                texto_memorial = conteudo_arquivo.decode("utf-8", errors="ignore")
        else:
            texto_memorial = conteudo_arquivo.decode("utf-8", errors="ignore")

        texto_memorial = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\xFF]', '', texto_memorial)
        return texto_memorial

    def gerar_documentos_pelo_memorial(
        self, conteudo_arquivo: bytes, nome_arquivo: str, dados_projeto: Dict[str, Any]
    ) -> List[Tuple[str, io.BytesIO]]:
        """
        Gera um documento Word separado para cada confrontante identificado no memorial.
        """
        if not conteudo_arquivo:
            raise ValueError("O conteúdo do arquivo está vazio ou corrompido.")

        texto_memorial = self._extrair_texto_memorial(conteudo_arquivo, nome_arquivo)
        dados_ia = self._obter_dados_estruturados_com_ia(texto_memorial, dados_projeto)
        
        # Recupera proprietário e cpf de origem extraídos pela IA do texto
        proprietario_origem = dados_ia.get("proprietario_origem", dados_projeto.get("proprietario", "RODRIGO COLOMBI FROTA")).upper()
        cpf_origem = dados_ia.get("cpf_origem", dados_projeto.get("cpf_proprietario", "092.653.737-76"))

        # Injeta esses dados em dados_projeto para passar para a montagem de cada documento
        dados_projeto_atualizados = dados_projeto.copy()
        dados_projeto_atualizados["proprietario"] = proprietario_origem
        dados_projeto_atualizados["cpf_proprietario"] = cpf_origem

        confrontantes = dados_ia.get("confrontantes") or self._estrutura_padrao()["confrontantes"]

        documentos: List[Tuple[str, io.BytesIO]] = []
        for dados_confrontante in confrontantes:
            nome_confrontante = str(
                dados_confrontante.get("confrontante_proprietario", "Confrontante")
            ).strip()
            buffer = self._montar_documento_confrontante(dados_confrontante, dados_projeto_atualizados)
            documentos.append((nome_confrontante, buffer))

        return documentos

    def gerar_documento_pelo_memorial(
        self, conteudo_arquivo: bytes, nome_arquivo: str, dados_projeto: Dict[str, Any]
    ) -> io.BytesIO:
        """
        Mantido para compatibilidade simples (retorna o primeiro).
        """
        documentos = self.gerar_documentos_pelo_memorial(conteudo_arquivo, nome_arquivo, dados_projeto)
        return documentos[0][1]

    @staticmethod
    def gerar_zip_anuencias(documentos: List[Tuple[str, io.BytesIO]], prefixo_arquivo: str = "ANUENCIA_INCRA") -> io.BytesIO:
        zip_buffer = io.BytesIO()
        nomes_usados = set()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for nome_confrontante, buffer in documentos:
                nome_base = re.sub(r'[^A-Za-z0-9_\-]+', '_', nome_confrontante.strip().upper()) or "CONFRONTANTE"
                nome_arquivo = f"{prefixo_arquivo}_{nome_base}.docx"
                sufixo = 2
                nome_final = nome_arquivo
                while nome_final in nomes_usados:
                    nome_final = f"{prefixo_arquivo}_{nome_base}_{sufixo}.docx"
                    sufixo += 1
                nomes_usados.add(nome_final)

                buffer.seek(0)
                zip_file.writestr(nome_final, buffer.read())
                buffer.seek(0)

        zip_buffer.seek(0)
        return zip_buffer

    def _definir_margens_celulas_zero(self, cell):
        """
        Remove o preenchimento interno padrão (padding) das células do Word.
        """
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m in ['top', 'bottom', 'left', 'right']:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), '20')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def _cm_para_dxa(self, valor_cm: float) -> int:
        """
        Converte centímetros para Twips (dxa), que é a unidade usada
        internamente pelo Word para larguras de tabela/coluna.
        1 cm = 566.9291339 dxa (1 polegada = 1440 dxa, 1 polegada = 2.54 cm).
        """
        return int(round(valor_cm * 1440 / 2.54))

    def _forcar_largura_fixa_tabela(self, tabela, larguras_cm: List[float]):
        """
        Garante que o Word respeite as larguras de coluna definidas.
        """
        tbl = tabela._tbl
        tblPr = tbl.tblPr

        # Centraliza a tabela na página
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')

        # 1. Trava o layout da tabela em "fixed"
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'fixed')

        # 2. Reconstrói o tblGrid com as larguras exatas
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            tbl.remove(tblGrid)
        tblGrid = OxmlElement('w:tblGrid')
        for largura_cm in larguras_cm:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(self._cm_para_dxa(largura_cm)))
            tblGrid.append(gridCol)
        # tblGrid deve vir logo após tblPr
        tblPr.addnext(tblGrid)

        # 3. Reaplica a largura em todas as linhas/células já existentes
        for row in tabela.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(larguras_cm):
                    cell.width = Cm(larguras_cm[idx])
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), str(self._cm_para_dxa(larguras_cm[idx])))
                    tcW.set(qn('w:type'), 'dxa')

    def _montar_documento_confrontante(
        self, dados_ia: Dict[str, Any], dados_projeto: Dict[str, Any]
    ) -> io.BytesIO:
        """
        Gera o documento idêntico ao modelo fornecido.
        """
        doc = Document()

        # Configurações de página (Paisagem com margens otimizadas)
        for section in doc.sections:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)

        # 1. TÍTULO
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run("DECLARAÇÃO DE RESPEITO DE LIMITES")
        run_titulo.bold = True
        run_titulo.font.size = Pt(12)
        run_titulo.font.name = 'Arial'
        run_titulo.font.color.rgb = RGBColor(0, 100, 0)

        # 2. TEXTO DA DECLARAÇÃO
        prop = dados_projeto.get("proprietario", "RODRIGO COLOMBI FROTA").upper()
        cpf = dados_projeto.get("cpf_proprietario", "092.653.737-76")
        tec_nome = self.dados_tecnico.get("nome", "Régis Campo da Silva")
        tec_cfta = self.dados_tecnico.get("cfta", "11198519711")

        p_corpo = doc.add_paragraph()
        p_corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_corpo.paragraph_format.space_before = Pt(12)
        p_corpo.add_run(f"Eu, {prop}, CPF {cpf}, residente em Vila Valério, e eu, {tec_nome}, Tecnico em Agropecuaria, CFTA {tec_cfta}, credenciado pelo INCRA sob o codigo G1D, declaramos sob as penas da Lei que quando dos trabalhos topograficos executados na citada propriedade foram respeitados os limites de \"divisas in loco\" com os confrontantes abaixo relacionados, não havendo qualquer litigio entre as partes.")

        # 3. LABEL CONFRONTANTES
        p_label = doc.add_paragraph()
        p_label.paragraph_format.space_before = Pt(12)
        run_label = p_label.add_run("Confrontantes:")
        run_label.bold = True

        # 4. TABELA DE VÉRTICES
        tabela = doc.add_table(rows=1, cols=8)
        tabela.style = 'Table Grid'
        tabela.autofit = False

        headers = ["Código", "Longitude", "Latitude", "Altitude (m)", "Código", "Azimute", "Dist. (m)", "Confrontante"]
        # Novas larguras exigidas (em cm), na ordem das colunas:
        # 1 Código | 2 Longitude | 3 Latitude | 4 Altitude | 5 Código | 6 Azimute | 7 Dist. | 8 Confrontante
        larguras_cm = [2.39, 2.5, 2.38, 2.25, 2.00, 1.75, 1.75, 8.75]
        larguras = [Cm(v) for v in larguras_cm]

        for idx, text in enumerate(headers):
            cell = tabela.rows[0].cells[idx]
            cell.width = larguras[idx]
            cell.text = text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            self._definir_margens_celulas_zero(cell)

        for v in dados_ia.get("vertices", []):
            row = tabela.add_row()
            vals = [
                v.get("codigo"), 
                self._formatar_coordenada(v.get("longitude")), 
                self._formatar_coordenada(v.get("latitude")), 
                v.get("altitude"), 
                v.get("vante"), 
                self._formatar_azimute(v.get("azimute")), 
                v.get("distancia"), 
                v.get("confrontacao_completa")
            ]
            for i in range(8):
                cell = row.cells[i]
                cell.width = larguras[i]
                cell.text = str(vals[i])
                self._definir_margens_celulas_zero(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centraliza o texto na célula
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if p.runs:
                    run = p.runs[0]
                    run.font.size = Pt(9) # Fonte da tabela tamanho 9
                    run.font.name = 'Arial Narrow'

        # Trava as larguras definitivamente e centraliza a tabela
        self._forcar_largura_fixa_tabela(tabela, larguras_cm)

        # 5. ASSINATURAS (Proprietário e Confrontante)
        doc.add_paragraph().paragraph_format.space_before = Pt(36)
        
        # Tabela para assinaturas do Proprietário e Confrontante
        tab_ass = doc.add_table(rows=2, cols=2)
        tab_ass.autofit = False
        tab_ass.columns[0].width = Cm(13.5)
        tab_ass.columns[1].width = Cm(13.5)
        
        # Linhas de assinatura
        p0 = tab_ass.rows[0].cells[0].paragraphs[0]
        p0.add_run("__________________________________________________")
        p1 = tab_ass.rows[0].cells[1].paragraphs[0]
        p1.add_run("__________________________________________________")
        
        # Nomes e CPFs
        p_nome0 = tab_ass.rows[1].cells[0].paragraphs[0]
        p_nome0.add_run(f"{prop}\n{cpf}")
        p_nome1 = tab_ass.rows[1].cells[1].paragraphs[0]
        p_nome1.add_run(f"{str(dados_ia.get('confrontante_proprietario', '')).upper()}")

        # 6. ANEXOS
        doc.add_paragraph().paragraph_format.space_before = Pt(24)
        p_anexos = doc.add_paragraph()
        run_anexos = p_anexos.add_run("Anexos: Planta do Imóvel Memorial Descritivo do Imóvel")
        run_anexos.font.size = Pt(9)

        # 7. DATA (Alinhada à direita)
        p_data = doc.add_paragraph()
        p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_data.paragraph_format.space_before = Pt(12)
        meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
        hoje = datetime.now()
        p_data.add_run(f"Vila Valério, {hoje.day} de {meses[hoje.month-1]} de {hoje.year}.")

        # 8. ASSINATURA DO TÉCNICO (Régis) - FINAL DO DOCUMENTO
        doc.add_paragraph().paragraph_format.space_before = Pt(48)
        p_ass_tec = doc.add_paragraph()
        p_ass_tec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ass_tec.add_run("__________________________________________________\n")
        run_tec = p_ass_tec.add_run(f"{tec_nome}\nTecnico em Agropecuaria - CFTA: {tec_cfta}")
        run_tec.bold = True
        run_tec.font.size = Pt(10)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
